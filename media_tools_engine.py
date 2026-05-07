"""
media_tools_engine.py
ExamHelp AI — Media Tools Suite
  Tab 1: AI Photo Editor (20+ effects)
  Tab 2: Image → PDF
  Tab 3: PDF → Any Format
  Tab 4: File QR Share (up to 10 MB)
"""
from __future__ import annotations
import base64, io, os, json, time, random, math, urllib.parse
from typing import List, Optional, Tuple
import streamlit as st
from PIL import Image, ImageFilter, ImageEnhance, ImageOps, ImageDraw

try:
    import numpy as np
    _NP = True
except ImportError:
    _NP = False

try:
    import fitz
    _FITZ = True
except ImportError:
    _FITZ = False

try:
    import qrcode
    _QR = True
except ImportError:
    _QR = False

try:
    import requests as _req
    _REQ = True
except ImportError:
    _REQ = False

try:
    from docx import Document as _DocxDoc
    from docx.shared import Inches, Pt
    _DOCX = True
except ImportError:
    _DOCX = False

try:
    from rembg import remove as _rembg_remove
    _REMBG = True
except ImportError:
    _REMBG = False


# ═══════════════════════════════════════════════════════════
# CSS
# ═══════════════════════════════════════════════════════════
_MT_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;700&family=Syne:wght@700;800&display=swap');

.mt-hero {
    background: radial-gradient(ellipse at 20% 0%, rgba(99,102,241,0.18) 0%, transparent 60%),
                radial-gradient(ellipse at 80% 100%, rgba(6,182,212,0.14) 0%, transparent 60%),
                rgba(10,14,30,0.85);
    border: 1px solid rgba(99,102,241,0.22);
    border-radius: 24px; padding: 36px 28px 32px;
    margin-bottom: 24px; text-align: center;
    backdrop-filter: blur(20px);
    position: relative; overflow: hidden;
    animation: mtFadeUp 0.6s cubic-bezier(0.16,1,0.3,1) both;
}
.mt-hero::before {
    content:''; position:absolute; top:-1px; left:10%; right:10%; height:2px;
    background: linear-gradient(90deg, transparent, #6366f1, #06b6d4, transparent);
    border-radius:100px;
}
.mt-hero-title {
    font-family:'Syne',sans-serif; font-size:clamp(1.8rem,3.5vw,2.6rem); font-weight:800;
    background: linear-gradient(135deg,#fff 0%,#c7d2fe 30%,#818cf8 55%,#06b6d4 80%);
    -webkit-background-clip:text; -webkit-text-fill-color:transparent; background-clip:text;
    margin-bottom:10px; filter:drop-shadow(0 0 24px rgba(99,102,241,0.3));
}
.mt-hero-sub {
    font-size:0.95rem; color:rgba(255,255,255,0.5); max-width:580px; margin:0 auto;
    line-height:1.7;
}
@keyframes mtFadeUp {
    from{opacity:0;transform:translateY(20px);} to{opacity:1;transform:none;}
}

/* Filter button grid */
.mt-filter-btn {
    background: rgba(10,14,30,0.7);
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: 12px; padding: 10px 8px;
    text-align: center; cursor: pointer;
    transition: all 0.25s ease;
    font-size: 0.75rem; color: rgba(255,255,255,0.7);
}
.mt-filter-btn:hover, .mt-filter-btn.active {
    border-color: rgba(99,102,241,0.5);
    background: rgba(99,102,241,0.1);
    color: #c7d2fe; transform: translateY(-2px);
    box-shadow: 0 8px 24px rgba(99,102,241,0.15);
}
.mt-filter-icon { font-size:1.4rem; display:block; margin-bottom:4px; }

/* Section header */
.mt-sec {
    display:flex; align-items:center; gap:10px; margin:20px 0 12px;
}
.mt-sec-line { flex:1; height:1px; background:linear-gradient(90deg,rgba(99,102,241,0.4),transparent); }
.mt-sec-label {
    font-family:'JetBrains Mono',monospace; font-size:0.62rem;
    letter-spacing:3px; color:#818cf8; text-transform:uppercase; white-space:nowrap;
}

/* Effect card */
.mt-fx-btn {
    position:relative; overflow:hidden;
    background:rgba(10,14,30,0.7); border:1px solid rgba(255,255,255,0.07);
    border-radius:14px; padding:14px 10px; text-align:center; cursor:pointer;
    transition:all 0.3s cubic-bezier(0.16,1,0.3,1);
}
.mt-fx-btn:hover { transform:translateY(-4px); border-color:rgba(99,102,241,0.45);
    box-shadow:0 12px 36px rgba(0,0,0,0.35); }
.mt-fx-icon { font-size:1.8rem; display:block; margin-bottom:6px; }
.mt-fx-label { font-size:0.7rem; font-weight:700; color:rgba(255,255,255,0.75); }

/* QR result card */
.mt-qr-card {
    background: linear-gradient(135deg,rgba(99,102,241,0.1),rgba(6,182,212,0.07));
    border: 1px solid rgba(99,102,241,0.25);
    border-radius: 20px; padding: 24px; text-align:center;
    animation: mtFadeUp 0.5s both;
}
.mt-qr-link {
    font-family:'JetBrains Mono',monospace; font-size:0.82rem;
    color:#a5f3fc; word-break:break-all; background:rgba(6,182,212,0.08);
    border:1px solid rgba(6,182,212,0.2); border-radius:10px;
    padding:10px 14px; display:block; margin:14px 0;
}
.mt-info-chip {
    display:inline-flex; align-items:center; gap:6px;
    background:rgba(255,255,255,0.04); border:1px solid rgba(255,255,255,0.08);
    border-radius:100px; padding:5px 14px; font-size:0.75rem;
    color:rgba(255,255,255,0.45); margin:4px;
}

/* Palette swatch */
.mt-swatch {
    display:inline-block; width:40px; height:40px;
    border-radius:8px; margin:4px;
    border:1px solid rgba(255,255,255,0.1);
    box-shadow:0 4px 12px rgba(0,0,0,0.3);
    transition:transform 0.2s ease;
}
.mt-swatch:hover { transform:scale(1.1); }

/* Stat tile */
.mt-stat { background:rgba(10,14,30,0.8); border:1px solid rgba(255,255,255,0.07);
    border-radius:14px; padding:14px; text-align:center; }
.mt-stat-n { font-family:'Syne',sans-serif; font-size:1.4rem; font-weight:800;
    color:#818cf8; }
.mt-stat-l { font-size:0.6rem; letter-spacing:2px; color:rgba(255,255,255,0.3);
    text-transform:uppercase; font-family:'JetBrains Mono',monospace; }
</style>
"""

# ═══════════════════════════════════════════════════════════
# PHOTO FILTER & EFFECT FUNCTIONS
# ═══════════════════════════════════════════════════════════

def _to_np(img):
    return np.array(img.convert("RGB"), dtype=np.float32)

def _from_np(arr):
    return Image.fromarray(np.clip(arr, 0, 255).astype(np.uint8))

def _apply_filter(img, name):
    img = img.convert("RGB")
    if name == "Original": return img
    if name == "Grayscale": return ImageOps.grayscale(img).convert("RGB")
    if name == "Sepia":
        if _NP:
            a=_to_np(img)
            return _from_np(np.stack([
                np.clip(a[:,:,0]*0.393+a[:,:,1]*0.769+a[:,:,2]*0.189,0,255),
                np.clip(a[:,:,0]*0.349+a[:,:,1]*0.686+a[:,:,2]*0.168,0,255),
                np.clip(a[:,:,0]*0.272+a[:,:,1]*0.534+a[:,:,2]*0.131,0,255)
            ],axis=2))
        return ImageEnhance.Color(img).enhance(0.3)
    if name == "Vintage":
        img=ImageEnhance.Color(ImageEnhance.Contrast(img).enhance(0.85)).enhance(0.7)
        if _NP:
            a=_to_np(img); a[:,:,0]=np.clip(a[:,:,0]*1.1,0,255); a[:,:,2]=np.clip(a[:,:,2]*0.8,0,255)
            return _from_np(a)
        return img
    if name == "Noir": return ImageOps.autocontrast(ImageOps.grayscale(img),10).convert("RGB")
    if name == "Vivid": return ImageEnhance.Contrast(ImageEnhance.Color(img).enhance(2.0)).enhance(1.2)
    if name == "Cool":
        if _NP:
            a=_to_np(img); a[:,:,0]=np.clip(a[:,:,0]*0.85,0,255); a[:,:,2]=np.clip(a[:,:,2]*1.15,0,255)
            return _from_np(a)
        return img
    if name == "Warm":
        if _NP:
            a=_to_np(img); a[:,:,0]=np.clip(a[:,:,0]*1.15,0,255); a[:,:,2]=np.clip(a[:,:,2]*0.85,0,255)
            return _from_np(a)
        return img
    if name == "Fade": return ImageEnhance.Contrast(ImageEnhance.Color(img).enhance(0.5)).enhance(0.75)
    if name == "Duotone":
        if _NP:
            gray=np.array(ImageOps.grayscale(img),dtype=np.float32)/255.0
            return _from_np(np.stack([gray*99,gray*102,gray*241],axis=2))
        return img
    if name == "Sunset":
        if _NP:
            a=_to_np(img); a[:,:,0]=np.clip(a[:,:,0]*1.2,0,255); a[:,:,2]=np.clip(a[:,:,2]*0.7,0,255)
            return _from_np(a)
        return img
    if name == "Forest":
        if _NP:
            a=_to_np(img); a[:,:,1]=np.clip(a[:,:,1]*1.15,0,255); a[:,:,0]=np.clip(a[:,:,0]*0.9,0,255)
            return _from_np(a)
        return img
    if name == "Cyberpunk":
        if _NP:
            a=_to_np(img)
            a[:,:,0]=np.clip(a[:,:,0]*1.3,0,255); a[:,:,2]=np.clip(a[:,:,2]*1.4,0,255); a[:,:,1]=np.clip(a[:,:,1]*0.7,0,255)
            return _from_np(a)
        return img
    if name == "Film":
        img=ImageEnhance.Color(ImageEnhance.Contrast(img).enhance(1.1)).enhance(0.85)
        if _NP:
            return _from_np(_to_np(img)+np.random.normal(0,6,np.array(img).shape).astype(np.float32))
        return img
    return img

def _fx_enhance(img): return ImageEnhance.Brightness(ImageEnhance.Color(ImageEnhance.Contrast(ImageEnhance.Sharpness(img).enhance(1.5)).enhance(1.15)).enhance(1.2)).enhance(1.05)
def _fx_hdr(img): return ImageEnhance.Color(ImageEnhance.Sharpness(ImageEnhance.Contrast(img).enhance(1.4)).enhance(2.0)).enhance(1.5)

def _fx_vignette(img):
    if not _NP: return img
    w,h=img.size; a=_to_np(img)
    X=np.linspace(-1,1,w); Y=np.linspace(-1,1,h)
    Xg,Yg=np.meshgrid(X,Y)
    return _from_np(a*np.clip(1.0-(Xg**2+Yg**2)*0.7,0.2,1.0)[:,:,np.newaxis])

def _fx_tiltshift(img):
    w,h=img.size; bl=img.filter(ImageFilter.GaussianBlur(8))
    mask=Image.new("L",(w,h),0); draw=ImageDraw.Draw(mask)
    cy,fh=h//2,h//5
    for y in range(h):
        d=abs(y-cy)
        v=255 if d<fh else max(0,int(255*(1-(d-fh)/fh))) if d<fh*2 else 0
        draw.line([(0,y),(w,y)],fill=v)
    return Image.composite(img,bl,mask)

def _fx_sketch(img):
    if not _NP: return ImageOps.grayscale(img).convert("RGB")
    gray=np.array(ImageOps.grayscale(img),dtype=np.float32)
    blur=np.array(Image.fromarray((255-gray).astype(np.uint8)).filter(ImageFilter.GaussianBlur(20)),dtype=np.float32)
    return Image.fromarray(np.clip(gray*255/(255-blur+1),0,255).astype(np.uint8)).convert("RGB")

def _fx_glitch(img):
    if not _NP: return img
    a=np.array(img.convert("RGB")); r=a.copy()
    r[:,:,0]=np.roll(a[:,:,0],random.randint(6,18),axis=1)
    for _ in range(random.randint(4,8)):
        y=random.randint(0,a.shape[0]-5)
        r[y:y+random.randint(2,6)]=np.roll(r[y:y+random.randint(2,6)],random.randint(-25,25),axis=1)
    return Image.fromarray(r)

def _fx_chromatic(img,s=6):
    if not _NP: return img
    a=np.array(img.convert("RGB")); r=a.copy()
    r[:,s:,0]=a[:,:-s,0]; r[:,:-s,2]=a[:,s:,2]
    return Image.fromarray(r)

def _fx_grain(img):
    if not _NP: return img
    return _from_np(_to_np(img)+np.random.normal(0,14,np.array(img).shape).astype(np.float32))

def _fx_dreamy(img):
    soft=img.filter(ImageFilter.GaussianBlur(3))
    if _NP: return _from_np(np.clip(_to_np(img)*0.6+_to_np(soft)*0.55,0,255))
    return soft

def _fx_neon(img):
    if not _NP: return img
    e=img.filter(ImageFilter.FIND_EDGES); ea=_to_np(e)
    ea[:,:,0]*=3; ea[:,:,1]*=2; ea[:,:,2]*=4
    return _from_np(np.clip(ea*2.5,0,255))

def _extract_palette(img,n=8):
    pixels=list(img.resize((80,80)).convert("RGB").getdata()); random.shuffle(pixels)
    colors=[]; seen=set()
    for r,g,b in pixels[:400]:
        hx=f"#{r:02x}{g:02x}{b:02x}"
        if hx not in seen: seen.add(hx); colors.append(hx)
        if len(colors)>=n: break
    return colors

def _img_to_bytes(img,fmt="PNG",quality=90):
    buf=io.BytesIO()
    img.convert("RGB").save(buf,format="JPEG" if fmt=="JPG" else fmt,**( {"quality":quality} if fmt in ("JPG","JPEG") else {}))
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════
# TAB 1 — AI PHOTO EDITOR
# ═══════════════════════════════════════════════════════════

_FILTERS = ["Original","Grayscale","Sepia","Vintage","Noir","Vivid",
            "Cool","Warm","Fade","Duotone","Sunset","Forest","Cyberpunk","Film"]

_FX = [
    ("✨","AI Enhance","enhance"),("🌅","HDR Effect","hdr"),
    ("🌑","Vignette","vignette"),("🔭","Tilt-Shift","tiltshift"),
    ("✏️","Sketch","sketch"),("⚡","Glitch","glitch"),
    ("🌈","Chromatic","chromatic"),("🎞️","Film Grain","grain"),
    ("💭","Dreamy","dreamy"),("🌟","Neon","neon"),
]

def _render_photo_editor():
    # ── Sub-tabs: AI Generator / Photo Editor / Unsplash Explorer ──
    st.markdown("""
    <div class="mt-sec">
        <span class="mt-sec-label">🎨 AI Studio</span>
        <div class="mt-sec-line"></div>
    </div>
    """, unsafe_allow_html=True)

    subtab1, subtab2, subtab3 = st.tabs(["🤖 AI Image Generator", "🖼️ Photo Editor", "🌄 Unsplash Explorer"])

    with subtab1:
        _render_ai_image_generator()

    with subtab2:
        _render_classic_photo_editor()

    with subtab3:
        _render_unsplash_explorer()


# ─────────────────────────────────────────────────────────────────────
# AI IMAGE GENERATOR  (Pollinations.ai — 100% free, no key)
# ─────────────────────────────────────────────────────────────────────

_GEN_STYLES = {
    "⭐ Flux Best":          ("flux",         "#6366f1", "Best all-round quality"),
    "📷 Photorealistic":     ("flux-realism",  "#0ea5e9", "Ultra-realistic photography"),
    "🎨 Artistic / Oil":     ("flux-cablyai",  "#f59e0b", "Painterly, gallery quality"),
    "🌸 Anime / Manga":      ("flux-anime",    "#ec4899", "Studio Ghibli style"),
    "💎 3D Render":          ("flux-3d",       "#06b6d4", "Octane/Blender quality"),
    "🌑 Dark / Gothic":      ("any-dark",      "#7c3aed", "Moody, dramatic"),
    "⚡ Turbo (Fast)":       ("turbo",         "#84cc16", "Fastest generation"),
    "🌈 DreamShaper":        ("dreamshaper",   "#f472b6", "Dreamy surreal art"),
}

_STYLE_SUFFIX = {
    "⭐ Flux Best":        "highly detailed, 8K, masterpiece quality",
    "📷 Photorealistic":   "ultra-photorealistic, DSLR, sharp focus, professional photography",
    "🎨 Artistic / Oil":   "oil painting, rich impasto, gallery quality, museum artwork",
    "🌸 Anime / Manga":    "anime art style, studio ghibli, vibrant illustration, cel-shaded",
    "💎 3D Render":        "3D render, octane, blender, PBR materials, cinematic lighting",
    "🌑 Dark / Gothic":    "dark gothic, dramatic shadows, chiaroscuro, moody atmosphere",
    "⚡ Turbo (Fast)":     "detailed, high quality",
    "🌈 DreamShaper":      "dreamlike surreal, soft colors, magical atmosphere, ethereal",
}

_ASPECT_SIZES = {
    "1:1 Square (1024×1024)":    (1024, 1024),
    "16:9 Landscape (1280×720)": (1280, 720),
    "9:16 Portrait (720×1280)":  (720,  1280),
    "4:3 Classic (1024×768)":    (1024, 768),
    "3:2 Photo (1200×800)":      (1200, 800),
    "21:9 Ultrawide (1344×576)": (1344, 576),
}

_PROMPT_EXAMPLES = [
    "A neon-lit cyberpunk city at night, flying cars, rain, hyper-detailed",
    "A serene Japanese garden with cherry blossoms falling into a koi pond",
    "Portrait of an astronaut floating in space, Earth visible behind visor",
    "Ancient dragon coiled around a mountain peak, stormy sky, epic scale",
    "Cozy coffee shop interior, warm lighting, steam rising from cups",
    "A futuristic library with holographic books, endless shelves",
    "Ocean sunset with golden light, sailboat silhouette, mirror reflection",
    "A forest spirit made of starlight and moss, glowing softly",
    "Hyperrealistic close-up of a bee on lavender flowers, bokeh",
    "Victorian steampunk inventor's workshop, brass gears, floating blueprints",
]


def _gen_pollinations_url(prompt: str, w: int, h: int, model: str, seed: int) -> str:
    import urllib.parse
    encoded = urllib.parse.quote(prompt, safe="")
    return (f"https://image.pollinations.ai/prompt/{encoded}"
            f"?width={w}&height={h}&model={model}&seed={seed}&nologo=true&enhance=false")


def _render_ai_image_generator():
    st.markdown("""
    <div style="background:linear-gradient(135deg,rgba(99,102,241,0.1),rgba(236,72,153,0.06));
        border:1px solid rgba(99,102,241,0.2);border-radius:16px;padding:16px 20px;margin-bottom:18px;">
        <div style="font-size:0.85rem;color:rgba(255,255,255,0.6);">
            ✨ <strong style="color:#c7d2fe;">8 Free AI Models</strong> · Powered by Pollinations.ai &amp; HuggingFace ·
            No API key needed · 100% free · HD quality
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Prompt area ──
    prompt = st.text_area(
        "Describe your image",
        placeholder="A majestic dragon soaring over snow-capped mountains at golden hour, 8K detail...",
        height=100, key="ai_gen_prompt", label_visibility="collapsed"
    )

    # Random prompt button
    c_rnd, c_clear = st.columns([1, 1])
    with c_rnd:
        if st.button("🎲 Random Prompt", key="ai_gen_rnd", use_container_width=True):
            st.session_state.ai_gen_prompt = random.choice(_PROMPT_EXAMPLES)
            st.rerun()
    with c_clear:
        if st.button("✏️ Clear Prompt", key="ai_gen_clear", use_container_width=True):
            st.session_state.ai_gen_prompt = ""
            st.rerun()

    # ── Style selector ──
    st.markdown("""<div class="mt-sec"><span class="mt-sec-label">🎨 Model / Style</span><div class="mt-sec-line"></div></div>""", unsafe_allow_html=True)

    style_cols = st.columns(4)
    chosen_style = st.session_state.get("ai_gen_style", "⭐ Flux Best")
    for i, (style_name, (model_id, color, desc)) in enumerate(_GEN_STYLES.items()):
        with style_cols[i % 4]:
            is_active = (chosen_style == style_name)
            border = f"border:2px solid {color};" if is_active else "border:1px solid rgba(255,255,255,0.07);"
            bg = f"background:rgba({','.join(str(int(color[j:j+2],16)) for j in (1,3,5))},0.12);" if is_active else "background:rgba(10,14,30,0.7);"
            st.markdown(f"""
            <div style="{bg}{border}border-radius:12px;padding:10px 8px;text-align:center;
                margin-bottom:8px;cursor:pointer;" title="{desc}">
                <div style="font-size:1.4rem;">{style_name.split()[0]}</div>
                <div style="font-size:0.65rem;color:rgba(255,255,255,0.55);margin-top:2px;font-weight:600;">{style_name.split(' ',1)[1]}</div>
                <div style="font-size:0.55rem;color:rgba(255,255,255,0.3);">{desc}</div>
            </div>""", unsafe_allow_html=True)
            if st.button(f"Select", key=f"ai_style_{i}", use_container_width=True,
                         type="primary" if is_active else "secondary"):
                st.session_state.ai_gen_style = style_name
                st.rerun()

    # ── Options row ──
    st.markdown("""<div class="mt-sec"><span class="mt-sec-label">⚙️ Settings</span><div class="mt-sec-line"></div></div>""", unsafe_allow_html=True)
    o1, o2, o3, o4 = st.columns(4)
    with o1:
        aspect = st.selectbox("Aspect Ratio", list(_ASPECT_SIZES.keys()), key="ai_gen_aspect")
    with o2:
        seed_mode = st.selectbox("Seed", ["Random", "Fixed"], key="ai_gen_seed_mode")
    with o3:
        fixed_seed = st.number_input("Seed value", 1, 999999, 42, key="ai_gen_fixed_seed",
                                      disabled=(seed_mode == "Random"))
    with o4:
        n_images = st.selectbox("How many?", [1, 2, 4, 6], key="ai_gen_n")

    negative = st.text_input("Negative prompt (optional)",
                              placeholder="blurry, low quality, watermark, text, deformed, ugly",
                              key="ai_gen_neg")

    # ── Generate button ──
    gen_col, _ = st.columns([2, 1])
    with gen_col:
        do_gen = st.button("🚀 Generate Images", type="primary", use_container_width=True, key="ai_gen_btn")

    if do_gen:
        if not prompt or not prompt.strip():
            st.warning("⚠️ Please enter a prompt first.")
            return

        style_name = st.session_state.get("ai_gen_style", "⭐ Flux Best")
        model_id, color, _ = _GEN_STYLES[style_name]
        w, h = _ASPECT_SIZES[aspect]
        suffix = _STYLE_SUFFIX[style_name]
        full_prompt = f"{prompt.strip()}, {suffix}"
        if negative:
            full_prompt += f" --no {negative}"

        seeds = []
        for _ in range(n_images):
            seeds.append(fixed_seed if seed_mode == "Fixed" else random.randint(1, 999999))

        st.session_state["ai_gen_results"] = []

        with st.spinner(f"🎨 Generating {n_images} image(s) with {style_name}..."):
            ncols = min(n_images, 2) if n_images > 1 else 1
            img_cols = st.columns(ncols)

            for idx, seed in enumerate(seeds):
                url = _gen_pollinations_url(full_prompt, w, h, model_id, seed)
                try:
                    import urllib.request as _ur
                    req = _ur.Request(url, headers={"User-Agent": "ExamHelpAI/2.0", "Accept": "image/*"})
                    with _ur.urlopen(req, timeout=60) as resp:
                        img_bytes = resp.read()
                    with img_cols[idx % ncols]:
                        st.image(img_bytes, caption=f"Seed {seed}", use_container_width=True)
                        st.download_button(f"⬇️ Save #{idx+1}", img_bytes,
                                           f"ai_gen_{idx+1}_seed{seed}.png", "image/png",
                                           use_container_width=True, key=f"ai_dl_{idx}_{seed}")
                    st.session_state["ai_gen_results"].append(img_bytes)
                except Exception as e:
                    with img_cols[idx % ncols]:
                        # Show embed URL as fallback
                        st.markdown(f"""
                        <div style="border-radius:12px;overflow:hidden;margin-bottom:8px;">
                            <img src="{url}" style="width:100%;border-radius:12px;" alt="Generated image"
                                onerror="this.src='https://via.placeholder.com/{w}x{h}?text=Loading...'"/>
                        </div>""", unsafe_allow_html=True)
                        st.caption(f"Seed: {seed} | [Open full size]({url})")

        st.success(f"✅ {n_images} image(s) generated! Prompt: _{prompt[:60]}..._")

    # ── Previously generated ──
    if st.session_state.get("ai_gen_results"):
        st.markdown("""<div class="mt-sec"><span class="mt-sec-label">📥 Last Batch</span><div class="mt-sec-line"></div></div>""", unsafe_allow_html=True)
        cols = st.columns(min(len(st.session_state["ai_gen_results"]), 4))
        for i, img_b in enumerate(st.session_state["ai_gen_results"]):
            with cols[i % len(cols)]:
                st.image(img_b, use_container_width=True)


# ─────────────────────────────────────────────────────────────────────
# UNSPLASH EXPLORER  (1000+ images via Unsplash API)
# ─────────────────────────────────────────────────────────────────────

_UNSPLASH_CATEGORIES = [
    "nature", "architecture", "technology", "people", "animals", "food",
    "travel", "abstract", "fashion", "sports", "business", "city",
    "art", "vintage", "minimal", "dark", "colorful", "ocean", "mountains",
    "forest", "flowers", "space", "cars", "music", "books",
]


def _unsplash_grid_url(query: str, page: int = 1, per_page: int = 30) -> list[dict]:
    """Fetch images from Unsplash Source API (no key, free, public CDN)."""
    import urllib.request, urllib.parse
    results = []
    # Use picsum + unsplash CDN approach for guaranteed images
    # Unsplash Source API (source.unsplash.com) provides random images by keyword
    q = urllib.parse.quote(query)
    for i in range(per_page):
        seed = (page - 1) * per_page + i + 1
        w, h = 400, 300
        # Use different seed-based approaches for variety
        url = f"https://picsum.photos/seed/{q}{seed}/{w}/{h}"
        unsplash_url = f"https://source.unsplash.com/{w}x{h}/?{q}&sig={seed}"
        results.append({
            "thumb": unsplash_url,
            "full": unsplash_url.replace(f"{w}x{h}", "1920x1080"),
            "id": f"{query}_{seed}",
            "author": "Unsplash",
            "download": unsplash_url,
        })
    return results


def _render_unsplash_explorer():
    st.markdown("""
    <div style="background:linear-gradient(135deg,rgba(6,182,212,0.08),rgba(99,102,241,0.05));
        border:1px solid rgba(6,182,212,0.2);border-radius:16px;padding:14px 20px;margin-bottom:16px;">
        <div style="font-size:0.85rem;color:rgba(255,255,255,0.6);">
            🌄 <strong style="color:#a5f3fc;">Unsplash Photo Library</strong> ·
            Discover beautiful free photos · Click to download full resolution
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Search + category
    s1, s2 = st.columns([3, 1])
    with s1:
        search_q = st.text_input("Search photos", placeholder="mountains, sunset, city...",
                                  key="unsp_query", label_visibility="collapsed")
    with s2:
        per_row = st.selectbox("Columns", [3, 4, 5, 6], index=1, key="unsp_cols")

    # Category chips
    st.markdown("""<div class="mt-sec"><span class="mt-sec-label">📁 Categories</span><div class="mt-sec-line"></div></div>""", unsafe_allow_html=True)
    cat_cols = st.columns(8)
    active_cat = st.session_state.get("unsp_cat", "nature")
    for i, cat in enumerate(_UNSPLASH_CATEGORIES):
        with cat_cols[i % 8]:
            is_active = (active_cat == cat)
            if st.button(cat, key=f"unsp_cat_{cat}", use_container_width=True,
                         type="primary" if is_active else "secondary"):
                st.session_state.unsp_cat = cat
                st.rerun()

    query = search_q.strip() if search_q.strip() else active_cat

    # Pagination — show 1000+ images across pages
    c_pg1, c_pg2, c_pg3 = st.columns([1, 2, 1])
    with c_pg2:
        page = st.slider("Page (30 images/page = 1000+ total)", 1, 34, 1, key="unsp_page")

    total_shown = page * 30
    st.markdown(f"""
    <div style="text-align:center;color:rgba(255,255,255,0.35);font-size:0.78rem;margin-bottom:12px;">
        Showing page {page} · up to {total_shown} images loaded · Query: <strong style="color:#818cf8;">{query}</strong>
    </div>""", unsafe_allow_html=True)

    # Render image grid
    n_cols = per_row
    count_per_page = 30
    img_data = _unsplash_grid_url(query, page, count_per_page)

    rows = [img_data[i:i+n_cols] for i in range(0, len(img_data), n_cols)]
    for row in rows:
        cols = st.columns(n_cols)
        for j, item in enumerate(row):
            with cols[j]:
                st.markdown(f"""
                <div style="border-radius:10px;overflow:hidden;margin-bottom:4px;
                    border:1px solid rgba(255,255,255,0.05);">
                    <a href="{item['full']}" target="_blank">
                        <img src="{item['thumb']}"
                            style="width:100%;height:160px;object-fit:cover;display:block;
                            transition:transform 0.3s ease;"
                            loading="lazy"
                            onerror="this.style.background='#1e293b';this.style.minHeight='120px';"
                        />
                    </a>
                </div>
                <div style="font-size:0.6rem;color:rgba(255,255,255,0.25);
                    text-align:center;margin-bottom:8px;">
                    📷 {item['author']} · <a href="{item['full']}" target="_blank"
                    style="color:#6366f1;text-decoration:none;">Full size ↗</a>
                </div>""", unsafe_allow_html=True)

    # Nav buttons
    b1, b2, b3 = st.columns(3)
    with b1:
        if page > 1 and st.button("← Previous page", key="unsp_prev", use_container_width=True):
            st.session_state.unsp_page = page - 1; st.rerun()
    with b2:
        st.markdown(f"<div style='text-align:center;color:rgba(255,255,255,0.4);padding-top:8px;'>Page {page} of 34</div>", unsafe_allow_html=True)
    with b3:
        if page < 34 and st.button("Next page →", key="unsp_next", use_container_width=True):
            st.session_state.unsp_page = page + 1; st.rerun()

    st.caption(f"📸 {count_per_page * 34:,}+ photos available across all pages · Powered by Unsplash")


# ─────────────────────────────────────────────────────────────────────
# CLASSIC PHOTO EDITOR  (original feature, now in sub-tab)
# ─────────────────────────────────────────────────────────────────────

def _render_classic_photo_editor():
    st.markdown("""<div class="mt-sec"><span class="mt-sec-label">🖼️ Photo Editor — Filters & Effects</span><div class="mt-sec-line"></div></div>""", unsafe_allow_html=True)
    up = st.file_uploader("Upload Image", type=["png","jpg","jpeg","webp","bmp"],
                          key="mt_photo_up", label_visibility="collapsed")
    if not up:
        st.markdown("""
        <div style="border:2px dashed rgba(99,102,241,0.3);border-radius:18px;padding:44px;text-align:center;">
            <div style="font-size:3rem;margin-bottom:12px;">🖼️</div>
            <div style="color:rgba(255,255,255,0.5);font-size:0.9rem;">Drop an image to start editing</div>
            <div style="color:rgba(255,255,255,0.25);font-size:0.75rem;margin-top:6px;">PNG · JPG · WEBP · BMP</div>
        </div>
        """, unsafe_allow_html=True)
        return

    orig = Image.open(up)
    ck = f"mt_orig_{up.name}_{up.size}"
    if st.session_state.get("mt_img_ck") != ck:
        st.session_state.mt_img_ck = ck
        st.session_state.mt_img = orig.copy()
        st.session_state.mt_filter = "Original"
        st.session_state.mt_brightness = 1.0
        st.session_state.mt_contrast = 1.0
        st.session_state.mt_saturation = 1.0
        st.session_state.mt_sharpness = 1.0
        st.session_state.mt_blur = 0

    # ── Show original + preview side-by-side ──
    c1, c2 = st.columns(2)
    with c1:
        st.caption("📷 Original")
        st.image(orig, use_container_width=True)
    with c2:
        st.caption("✨ Preview")
        preview_ph = st.empty()

    # ── Filter Gallery ──
    st.markdown("""<div class="mt-sec"><span class="mt-sec-label">🎨 Filters</span><div class="mt-sec-line"></div></div>""", unsafe_allow_html=True)
    fcols = st.columns(7)
    chosen_filter = st.session_state.get("mt_filter","Original")
    for i, f in enumerate(_FILTERS):
        with fcols[i % 7]:
            if st.button(f, key=f"mt_f_{f}", use_container_width=True,
                         type="primary" if chosen_filter==f else "secondary"):
                st.session_state.mt_filter = f
                st.rerun()

    # ── Adjustments ──
    st.markdown("""<div class="mt-sec"><span class="mt-sec-label">⚙️ Adjustments</span><div class="mt-sec-line"></div></div>""", unsafe_allow_html=True)
    a1,a2,a3,a4,a5 = st.columns(5)
    with a1: br = st.slider("☀️ Brightness", 0.1, 3.0, st.session_state.get("mt_brightness",1.0), 0.05, key="mt_br")
    with a2: co = st.slider("🎭 Contrast",   0.1, 3.0, st.session_state.get("mt_contrast",1.0),   0.05, key="mt_co")
    with a3: sa = st.slider("🌈 Saturation", 0.0, 3.0, st.session_state.get("mt_saturation",1.0), 0.05, key="mt_sa")
    with a4: sh = st.slider("🔪 Sharpness",  0.0, 5.0, st.session_state.get("mt_sharpness",1.0),  0.1,  key="mt_sh")
    with a5: bl = st.slider("💨 Blur",        0,   12,  st.session_state.get("mt_blur",0),          1,    key="mt_bl")

    # ── AI Effects ──
    st.markdown("""<div class="mt-sec"><span class="mt-sec-label">⚡ AI & Creative Effects</span><div class="mt-sec-line"></div></div>""", unsafe_allow_html=True)
    ec = st.columns(5)
    chosen_fx = None
    for i,(icon,label,fid) in enumerate(_FX):
        with ec[i%5]:
            if st.button(f"{icon} {label}", key=f"mt_fx_{fid}", use_container_width=True):
                chosen_fx = fid

    # ── Transforms ──
    st.markdown("""<div class="mt-sec"><span class="mt-sec-label">🔄 Transform</span><div class="mt-sec-line"></div></div>""", unsafe_allow_html=True)
    t1,t2,t3,t4,t5 = st.columns(5)
    do_rot90 = t1.button("↻ 90°",  key="mt_rot90",  use_container_width=True)
    do_rot180= t2.button("↺ 180°", key="mt_rot180", use_container_width=True)
    do_fh    = t3.button("↔ Flip H",key="mt_fh",    use_container_width=True)
    do_fv    = t4.button("↕ Flip V",key="mt_fv",    use_container_width=True)
    do_mirror= t5.button("🪞 Mirror",key="mt_mirror",use_container_width=True)

    wr1,wr2 = st.columns(2)
    with wr1:
        nw = st.number_input("Width px",  value=orig.width,  min_value=50, max_value=4000, step=10, key="mt_nw")
    with wr2:
        nh = st.number_input("Height px", value=orig.height, min_value=50, max_value=4000, step=10, key="mt_nh")
    do_resize = st.button("📐 Apply Resize", key="mt_resize", use_container_width=True)

    # ── Watermark ──
    st.markdown("""<div class="mt-sec"><span class="mt-sec-label">🖊️ Watermark</span><div class="mt-sec-line"></div></div>""", unsafe_allow_html=True)
    wm1,wm2,wm3 = st.columns(3)
    wm_text  = wm1.text_input("Watermark text", placeholder="© ExamHelp AI", key="mt_wm_txt")
    wm_pos   = wm2.selectbox("Position", ["Bottom-Right","Bottom-Left","Top-Right","Top-Left","Center"], key="mt_wm_pos")
    wm_alpha = wm3.slider("Opacity", 0.1, 1.0, 0.5, 0.05, key="mt_wm_alpha")
    do_wm = st.button("💧 Add Watermark", key="mt_add_wm", use_container_width=True)

    # ── Color Palette ──
    do_palette = st.button("🎨 Extract Color Palette", key="mt_palette", use_container_width=True)

    # ── Build edited image ──
    edited = st.session_state.mt_img.copy()
    edited = _apply_filter(edited, st.session_state.mt_filter)
    if br != 1.0: edited = ImageEnhance.Brightness(edited).enhance(br)
    if co != 1.0: edited = ImageEnhance.Contrast(edited).enhance(co)
    if sa != 1.0: edited = ImageEnhance.Color(edited).enhance(sa)
    if sh != 1.0: edited = ImageEnhance.Sharpness(edited).enhance(sh)
    if bl > 0:    edited = edited.filter(ImageFilter.GaussianBlur(radius=bl))

    # Apply AI effect
    FX_MAP = {"enhance":_fx_enhance,"hdr":_fx_hdr,"vignette":_fx_vignette,
              "tiltshift":_fx_tiltshift,"sketch":_fx_sketch,"glitch":_fx_glitch,
              "chromatic":_fx_chromatic,"grain":_fx_grain,"dreamy":_fx_dreamy,"neon":_fx_neon}
    if chosen_fx and chosen_fx in FX_MAP:
        with st.spinner(f"Applying {chosen_fx}..."):
            edited = FX_MAP[chosen_fx](edited)
        st.session_state.mt_img = edited

    # Apply transforms
    if do_rot90:   st.session_state.mt_img = edited.rotate(-90, expand=True); st.rerun()
    if do_rot180:  st.session_state.mt_img = edited.rotate(180);              st.rerun()
    if do_fh:      st.session_state.mt_img = ImageOps.mirror(edited);         st.rerun()
    if do_fv:      st.session_state.mt_img = ImageOps.flip(edited);           st.rerun()
    if do_mirror:  st.session_state.mt_img = ImageOps.mirror(edited);         st.rerun()
    if do_resize:  st.session_state.mt_img = edited.resize((int(nw),int(nh)), Image.LANCZOS); st.rerun()

    # Watermark
    if do_wm and wm_text:
        e2 = edited.convert("RGBA")
        overlay = Image.new("RGBA", e2.size, (255,255,255,0))
        draw = ImageDraw.Draw(overlay)
        tw, th = e2.size[0]//6, e2.size[1]//30
        pos_map = {
            "Bottom-Right": (e2.size[0]-tw*len(wm_text)-10, e2.size[1]-th*4),
            "Bottom-Left":  (10, e2.size[1]-th*4),
            "Top-Right":    (e2.size[0]-tw*len(wm_text)-10, 10),
            "Top-Left":     (10, 10),
            "Center":       (e2.size[0]//2-tw*len(wm_text)//2, e2.size[1]//2),
        }
        xy = pos_map.get(wm_pos, (10,10))
        draw.text(xy, wm_text, fill=(255,255,255,int(255*wm_alpha)))
        edited = Image.alpha_composite(e2, overlay).convert("RGB")
        st.session_state.mt_img = edited
        st.rerun()

    # Color Palette
    if do_palette:
        palette = _extract_palette(edited)
        swatches = "".join(f'<span class="mt-swatch" style="background:{c};" title="{c}"></span>' for c in palette)
        st.markdown(f"""
        <div style="background:rgba(10,14,30,0.8);border:1px solid rgba(255,255,255,0.07);
            border-radius:16px;padding:20px;margin-top:12px;">
            <div class="mt-sec-label" style="margin-bottom:12px;">🎨 Extracted Palette</div>
            <div>{swatches}</div>
            <div style="margin-top:12px;font-family:'JetBrains Mono',monospace;font-size:0.72rem;color:rgba(255,255,255,0.3);">
                {" · ".join(palette)}
            </div>
        </div>
        """, unsafe_allow_html=True)

    # Show preview
    preview_ph.image(edited, use_container_width=True)

    # ── Export ──
    st.markdown("""<div class="mt-sec"><span class="mt-sec-label">📥 Export</span><div class="mt-sec-line"></div></div>""", unsafe_allow_html=True)
    ex1,ex2,ex3,ex4 = st.columns(4)
    fmt_map = {"PNG":"PNG","JPG":"JPG","WEBP":"WEBP","BMP":"BMP"}
    for col,(fmt,fmtk) in zip([ex1,ex2,ex3,ex4],fmt_map.items()):
        with col:
            st.download_button(f"⬇️ {fmt}", _img_to_bytes(edited,fmtk),
                               f"edited.{fmt.lower()}", f"image/{fmt.lower()}",
                               use_container_width=True, key=f"mt_dl_{fmt}")

# ═══════════════════════════════════════════════════════════
# TAB 2 — IMAGE → PDF
# ═══════════════════════════════════════════════════════════

def _render_image_to_pdf():
    st.markdown("""<div class="mt-sec"><span class="mt-sec-label">🖼️ Image → PDF Converter</span><div class="mt-sec-line"></div></div>""", unsafe_allow_html=True)

    st.markdown("""
    <div style="background:linear-gradient(135deg,rgba(6,182,212,0.08),rgba(99,102,241,0.05));
        border:1px solid rgba(6,182,212,0.2);border-radius:14px;padding:12px 18px;margin-bottom:16px;">
        <div style="font-size:0.82rem;color:rgba(255,255,255,0.55);">
            📌 <strong style="color:#a5f3fc;">Tips:</strong>
            Upload multiple images — they'll appear in PDF in upload order.
            Supports PNG · JPG · WEBP · BMP · up to 2GB per file.
        </div>
    </div>
    """, unsafe_allow_html=True)

    files = st.file_uploader("Upload Images (multiple OK)", type=["png","jpg","jpeg","webp","bmp"],
                              accept_multiple_files=True, key="mt_img2pdf_up", label_visibility="collapsed")
    if not files:
        st.markdown("""
        <div style="border:2px dashed rgba(6,182,212,0.3);border-radius:18px;padding:40px;text-align:center;">
            <div style="font-size:3rem;margin-bottom:12px;">🖼️→📄</div>
            <div style="color:rgba(255,255,255,0.5);">Upload one or more images to convert to PDF</div>
            <div style="color:rgba(255,255,255,0.3);font-size:0.75rem;margin-top:8px;">
                Multiple files → multi-page PDF · Auto-sorted by filename
            </div>
        </div>""", unsafe_allow_html=True)
        return

    # ── File list with sizes ──
    total_kb = sum(f.size for f in files) / 1024
    st.markdown(f"""
    <div style="background:rgba(6,182,212,0.06);border:1px solid rgba(6,182,212,0.15);
        border-radius:14px;padding:14px 16px;margin-bottom:14px;">
        <div style="font-family:'JetBrains Mono',monospace;font-size:0.65rem;letter-spacing:2px;
            color:#06b6d4;margin-bottom:10px;">✅ {len(files)} IMAGE(S) LOADED · {total_kb:.1f} KB TOTAL</div>
        <div style="display:grid;grid-template-columns:repeat(auto-fill,minmax(240px,1fr));gap:6px;">
            {''.join(f"""<div style="display:flex;align-items:center;gap:8px;background:rgba(255,255,255,0.03);
                border-radius:8px;padding:6px 10px;">
                <span style="color:#06b6d4;">📄</span>
                <span style="color:rgba(255,255,255,0.75);font-size:0.8rem;flex:1;overflow:hidden;
                    text-overflow:ellipsis;white-space:nowrap;">{f.name}</span>
                <span style="color:rgba(255,255,255,0.3);font-size:0.7rem;">{f.size/1024:.1f}KB</span>
                <span style="background:rgba(6,182,212,0.15);color:#a5f3fc;border-radius:4px;
                    padding:1px 6px;font-size:0.6rem;">#{i+1}</span>
            </div>""" for i, f in enumerate(files))}
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Preview thumbnails
    thumb_cols = st.columns(min(len(files), 6))
    images = []
    for i, f in enumerate(files):
        img = Image.open(f).convert("RGB")
        images.append(img)
        with thumb_cols[i % 6]:
            st.image(img, caption=f"#{i+1} {f.name[:12]}", use_container_width=True)

    # Options
    c1,c2,c3 = st.columns(3)
    page_size = c1.selectbox("Page Size", ["A4","Letter","A3","A5","Same as image"], key="mt_i2p_size")
    orient    = c2.selectbox("Orientation", ["Portrait","Landscape"], key="mt_i2p_orient")
    quality   = c3.slider("JPEG Quality", 50, 100, 85, key="mt_i2p_q")
    fit_mode  = st.selectbox("Fit Mode", ["Fit to page (keep ratio)","Stretch to fill","Original size"], key="mt_i2p_fit")

    c_btn, c_info = st.columns([2, 1])
    with c_info:
        st.markdown(f"""
        <div style="background:rgba(10,14,30,0.7);border:1px solid rgba(255,255,255,0.06);
            border-radius:12px;padding:12px;text-align:center;margin-top:4px;">
            <div style="font-size:1.4rem;font-weight:800;color:#06b6d4;font-family:'Syne',sans-serif;">{len(images)}</div>
            <div style="font-size:0.6rem;color:rgba(255,255,255,0.3);letter-spacing:2px;text-transform:uppercase;">Pages in PDF</div>
        </div>""", unsafe_allow_html=True)

    with c_btn:
        do_convert = st.button("📄 Convert to PDF", type="primary", use_container_width=True, key="mt_i2p_btn")

    if do_convert:
        with st.spinner("Converting images to PDF..."):
            buf = io.BytesIO()
            PAGE_SIZES = {"A4":(595,842),"Letter":(612,792),"A3":(842,1191),"A5":(420,595)}

            if page_size == "Same as image":
                pw,ph = images[0].size
            else:
                pw,ph = PAGE_SIZES.get(page_size,(595,842))
                if orient == "Landscape": pw,ph = ph,pw

            processed = []
            prog = st.progress(0, "Processing images...")
            for idx, img in enumerate(images):
                prog.progress(int((idx+1)/len(images)*80), f"Processing image {idx+1}/{len(images)}...")
                img = img.convert("RGB")
                if page_size != "Same as image":
                    if fit_mode.startswith("Fit"):
                        img.thumbnail((pw,ph), Image.LANCZOS)
                        bg = Image.new("RGB",(pw,ph),(255,255,255))
                        offset=((pw-img.width)//2,(ph-img.height)//2)
                        bg.paste(img, offset)
                        img = bg
                    elif fit_mode.startswith("Stretch"):
                        img = img.resize((pw,ph), Image.LANCZOS)
                processed.append(img)

            prog.progress(90, "Building PDF...")
            processed[0].save(buf, format="PDF", save_all=True,
                              append_images=processed[1:], quality=quality)
            pdf_bytes = buf.getvalue()
            prog.progress(100, "Done!")
            time.sleep(0.3); prog.empty()

        st.success(f"✅ PDF created! ({len(pdf_bytes)//1024} KB · {len(images)} pages · {page_size} {orient})")
        st.download_button("⬇️ Download PDF", pdf_bytes, "images_converted.pdf",
                           "application/pdf", use_container_width=True, key="mt_i2p_dl")


# ═══════════════════════════════════════════════════════════
# TAB 3 — PDF → ANY FORMAT
# ═══════════════════════════════════════════════════════════

def _render_pdf_converter():
    st.markdown("""<div class="mt-sec"><span class="mt-sec-label">📄 PDF → Any Format</span><div class="mt-sec-line"></div></div>""", unsafe_allow_html=True)

    pdf_file = st.file_uploader("Upload PDF", type=["pdf"], key="mt_pdf_conv_up", label_visibility="collapsed")
    if not pdf_file:
        st.markdown("""
        <div style="border:2px dashed rgba(139,92,246,0.3);border-radius:18px;padding:40px;text-align:center;">
            <div style="font-size:3rem;margin-bottom:12px;">📄→🔀</div>
            <div style="color:rgba(255,255,255,0.5);">Upload a PDF to convert to another format</div>
        </div>""", unsafe_allow_html=True)
        return

    fmt = st.selectbox("Convert to:", ["Images (PNG)","Images (JPG)","Plain Text (.txt)",
                                        "Word Document (.docx)","HTML","Markdown (.md)"],
                        key="mt_pdf_fmt")
    dpi = st.slider("Resolution (DPI) — for image exports", 72, 300, 150, key="mt_pdf_dpi") if "Image" in fmt else 150

    if not st.button("🔀 Convert", type="primary", use_container_width=True, key="mt_pdf_conv_btn"):
        return

    if not _FITZ:
        st.error("PyMuPDF not installed. Run: pip install pymupdf"); return

    raw = pdf_file.read()
    doc = fitz.open(stream=raw, filetype="pdf")
    pg_count = doc.page_count

    with st.spinner(f"Converting {pg_count} page(s)..."):

        if "Images (PNG)" in fmt or "Images (JPG)" in fmt:
            ext = "jpg" if "JPG" in fmt else "png"
            import zipfile
            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for i, page in enumerate(doc):
                    mat = fitz.Matrix(dpi/72, dpi/72)
                    pix = page.get_pixmap(matrix=mat)
                    img_bytes = pix.tobytes(output=ext)
                    zf.writestr(f"page_{i+1:03d}.{ext}", img_bytes)
            zip_buf.seek(0)
            st.success(f"✅ {pg_count} pages converted!")
            st.download_button(f"⬇️ Download ZIP ({ext.upper()}s)", zip_buf.getvalue(),
                               f"pdf_pages.zip","application/zip", use_container_width=True, key="mt_pdf_dl_img")

        elif "Text" in fmt:
            txt = "\n\n".join(f"--- Page {i+1} ---\n{doc[i].get_text()}" for i in range(pg_count))
            st.success(f"✅ {len(txt.split())} words extracted!")
            st.text_area("Preview (first 2000 chars)", txt[:2000], height=200, key="mt_pdf_txt_prev")
            st.download_button("⬇️ Download TXT", txt.encode(), "pdf_text.txt",
                               "text/plain", use_container_width=True, key="mt_pdf_dl_txt")

        elif "Word" in fmt:
            if not _DOCX:
                st.error("python-docx not installed. Run: pip install python-docx"); return
            doc_word = _DocxDoc()
            doc_word.add_heading(pdf_file.name.replace(".pdf",""), 0)
            for i in range(pg_count):
                doc_word.add_heading(f"Page {i+1}", level=2)
                for para in doc[i].get_text().split("\n"):
                    if para.strip():
                        doc_word.add_paragraph(para.strip())
                doc_word.add_page_break()
            wbuf = io.BytesIO(); doc_word.save(wbuf)
            st.success("✅ Word document created!")
            st.download_button("⬇️ Download DOCX", wbuf.getvalue(), "pdf_converted.docx",
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               use_container_width=True, key="mt_pdf_dl_docx")

        elif "HTML" in fmt:
            pages_html = []
            for i in range(pg_count):
                pages_html.append(f"<section><h2>Page {i+1}</h2>{doc[i].get_text('html')}</section>")
            html_out = f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{pdf_file.name}</title></head><body>{''.join(pages_html)}</body></html>"
            st.success("✅ HTML created!")
            st.download_button("⬇️ Download HTML", html_out.encode(), "pdf_converted.html",
                               "text/html", use_container_width=True, key="mt_pdf_dl_html")

        elif "Markdown" in fmt:
            lines = []
            for i in range(pg_count):
                lines.append(f"\n## Page {i+1}\n")
                for ln in doc[i].get_text().split("\n"):
                    lines.append(ln)
            md_out = "\n".join(lines)
            st.success("✅ Markdown created!")
            st.download_button("⬇️ Download MD", md_out.encode(), "pdf_converted.md",
                               "text/markdown", use_container_width=True, key="mt_pdf_dl_md")

# ═══════════════════════════════════════════════════════════
# TAB 4 — FILE QR SHARE (upload → multi-backend → QR code)
# ═══════════════════════════════════════════════════════════

MAX_FILE_BYTES = 25 * 1024 * 1024  # 25 MB

def _upload_fileio(file_bytes: bytes, filename: str, expiry: str = "14d") -> dict:
    """Upload file to file.io and return response dict (fixed JSON parsing)."""
    import urllib.request, urllib.parse
    boundary = f"----FormBoundary{random.randint(100000,999999)}"
    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="{filename}"\r\n'
        f"Content-Type: application/octet-stream\r\n\r\n"
    ).encode() + file_bytes + f"\r\n--{boundary}--\r\n".encode()
    url = f"https://file.io/?expires={expiry}&autoDelete=true"
    req = urllib.request.Request(url, data=body,
          headers={"Content-Type": f"multipart/form-data; boundary={boundary}",
                   "User-Agent": "Mozilla/5.0 ExamHelp/2.0",
                   "Accept": "application/json"}, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            raw = resp.read().decode("utf-8", errors="ignore").strip()
            # Guard: file.io sometimes returns HTML on error
            if raw.startswith("<"):
                return {"success": False, "error": "Service returned HTML (may be temporarily down)"}
            data = json.loads(raw)
            # file.io v2 API uses "link" field
            link = data.get("link") or data.get("url") or data.get("key", "")
            if link and not link.startswith("http"):
                link = f"https://file.io/{link}"
            if data.get("success") or link:
                return {"success": True, "link": link}
            return {"success": False, "error": data.get("message", data.get("error", "Unknown error"))}
    except Exception as e:
        return {"success": False, "error": str(e)}


def _upload_transfer_sh(file_bytes: bytes, filename: str) -> dict:
    """Upload to transfer.sh (free, no account, files expire after 14 days)."""
    import urllib.request
    safe_name = urllib.parse.quote(filename, safe="")
    req = urllib.request.Request(
        f"https://transfer.sh/{safe_name}",
        data=file_bytes,
        headers={
            "Content-Type": "application/octet-stream",
            "Max-Days": "14",
            "Max-Downloads": "100",
            "User-Agent": "Mozilla/5.0 ExamHelp/2.0",
        },
        method="PUT",
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            link = resp.read().decode().strip()
            if link.startswith("https://"):
                return {"success": True, "link": link}
            return {"success": False, "error": "Unexpected response"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def _upload_tmpfiles(file_bytes: bytes, filename: str) -> dict:
    """Upload to tmpfiles.org (free, 1GB limit, expires in 60 days)."""
    import urllib.request
    boundary = f"----TmpFilesBoundary{random.randint(100000,999999)}"
    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="{filename}"\r\n'
        f"Content-Type: application/octet-stream\r\n\r\n"
    ).encode() + file_bytes + f"\r\n--{boundary}--\r\n".encode()
    req = urllib.request.Request(
        "https://tmpfiles.org/api/v1/upload",
        data=body,
        headers={
            "Content-Type": f"multipart/form-data; boundary={boundary}",
            "User-Agent": "Mozilla/5.0 ExamHelp/2.0",
            "Accept": "application/json",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            raw = resp.read().decode("utf-8", errors="ignore").strip()
            if raw.startswith("<"):
                return {"success": False, "error": "HTML response"}
            data = json.loads(raw)
            # tmpfiles returns {"status":"success","data":{"url":"..."}}
            link = (data.get("data") or {}).get("url", "")
            if link:
                # Convert to direct download link
                direct = link.replace("tmpfiles.org/", "tmpfiles.org/dl/")
                return {"success": True, "link": direct}
            return {"success": False, "error": data.get("message", "No URL")}
    except Exception as e:
        return {"success": False, "error": str(e)}


def _upload_file_smart(file_bytes: bytes, filename: str, expiry: str = "14d") -> dict:
    """Try multiple upload backends in order, return first success."""
    backends = [
        ("file.io", lambda: _upload_fileio(file_bytes, filename, expiry)),
        ("transfer.sh", lambda: _upload_transfer_sh(file_bytes, filename)),
        ("tmpfiles.org", lambda: _upload_tmpfiles(file_bytes, filename)),
    ]
    errors = []
    for name, fn in backends:
        try:
            result = fn()
            if result.get("success") and result.get("link"):
                result["backend"] = name
                return result
            errors.append(f"{name}: {result.get('error','no link')}")
        except Exception as e:
            errors.append(f"{name}: {e}")
    return {"success": False, "error": " | ".join(errors)}


def _make_qr(url: str) -> bytes:
    """Generate a styled QR code PNG."""
    if not _QR:
        return b""
    qr = qrcode.QRCode(error_correction=qrcode.constants.ERROR_CORRECT_H,
                        box_size=10, border=4)
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="#6366f1", back_color="white")
    buf = io.BytesIO(); img.save(buf, format="PNG")
    return buf.getvalue()


def _render_file_qr_share():
    st.markdown("""<div class="mt-sec"><span class="mt-sec-label">🔗 File QR Share — Upload & Get Link</span><div class="mt-sec-line"></div></div>""", unsafe_allow_html=True)

    st.markdown("""
    <div style="background:linear-gradient(135deg,rgba(99,102,241,0.08),rgba(6,182,212,0.05));
        border:1px solid rgba(99,102,241,0.2);border-radius:16px;padding:16px 20px;margin-bottom:20px;">
        <div style="display:flex;gap:16px;flex-wrap:wrap;">
            <span class="mt-info-chip">📁 Any file type</span>
            <span class="mt-info-chip">⚖️ Up to 25 MB</span>
            <span class="mt-info-chip">⏳ Auto-expiry</span>
            <span class="mt-info-chip">🔄 3 backup services</span>
            <span class="mt-info-chip">🆓 100% Free</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

    up = st.file_uploader("Upload any file (max 25 MB)", key="mt_qr_file", label_visibility="collapsed")
    c_exp, c_qr_size = st.columns(2)
    with c_exp:
        expiry = st.selectbox("Link expires after:", ["1d","3d","7d","14d","1w","1m"], index=3, key="mt_qr_expiry")
    with c_qr_size:
        qr_color = st.selectbox("QR Code style:", ["Indigo (#6366f1)","Cyan (#06b6d4)","Green (#10b981)","Pink (#ec4899)","Black (#000000)"], key="mt_qr_color")

    if not up:
        st.markdown("""
        <div style="border:2px dashed rgba(99,102,241,0.3);border-radius:18px;padding:44px;text-align:center;">
            <div style="font-size:3.5rem;margin-bottom:14px;">📁</div>
            <div style="color:rgba(255,255,255,0.5);font-size:0.9rem;">Upload any file up to <strong style="color:#818cf8;">25 MB</strong></div>
            <div style="color:rgba(255,255,255,0.25);font-size:0.75rem;margin-top:8px;">PDF · DOCX · ZIP · MP4 · PNG · Any format</div>
            <div style="margin-top:20px;color:rgba(255,255,255,0.2);font-size:0.72rem;">
                🔄 Automatically tries file.io → transfer.sh → tmpfiles.org
            </div>
        </div>""", unsafe_allow_html=True)
        return

    # Size check
    file_bytes = up.read()
    size_mb = len(file_bytes) / (1024 * 1024)

    # Show file info
    c1,c2,c3,c4 = st.columns(4)
    c1.markdown(f'<div class="mt-stat"><div class="mt-stat-n">{size_mb:.2f}</div><div class="mt-stat-l">MB</div></div>', unsafe_allow_html=True)
    c2.markdown(f'<div class="mt-stat"><div class="mt-stat-n">{up.name.split(".")[-1].upper()}</div><div class="mt-stat-l">Type</div></div>', unsafe_allow_html=True)
    c3.markdown(f'<div class="mt-stat"><div class="mt-stat-n">{len(file_bytes)//1024}</div><div class="mt-stat-l">KB</div></div>', unsafe_allow_html=True)
    c4.markdown(f'<div class="mt-stat"><div class="mt-stat-n">{up.name[:10]}…</div><div class="mt-stat-l">Name</div></div>', unsafe_allow_html=True)

    if len(file_bytes) > MAX_FILE_BYTES:
        st.error(f"❌ File is {size_mb:.2f} MB. Maximum allowed is 25 MB.")
        return

    if st.button("🚀 Upload & Generate QR Code", type="primary", use_container_width=True, key="mt_qr_upload_btn"):
        prog = st.progress(0, "Connecting to upload service...")
        status_ph = st.empty()

        for pct in range(0, 40, 8):
            time.sleep(0.08); prog.progress(pct, f"Preparing... {pct}%")

        status_ph.info("⬆️ Uploading — trying file.io first, then fallbacks automatically...")
        result = _upload_file_smart(file_bytes, up.name, expiry)

        if not result.get("success"):
            prog.empty(); status_ph.empty()
            st.error(f"❌ All upload services failed:\n{result.get('error','Unknown error')}")
            st.info("💡 Tips: check your internet connection, reduce file size, or try again in a moment.")
            return

        prog.progress(80, "Generating QR code...")
        status_ph.empty()
        link = result.get("link","")
        backend = result.get("backend", "cloud")

        # QR color extraction
        color_hex = qr_color.split("(")[1].rstrip(")")
        if _QR:
            qr = qrcode.QRCode(error_correction=qrcode.constants.ERROR_CORRECT_H, box_size=12, border=4)
            qr.add_data(link)
            qr.make(fit=True)
            qr_img = qr.make_image(fill_color=color_hex, back_color="white")
            qr_buf = io.BytesIO(); qr_img.save(qr_buf, format="PNG")
            qr_bytes = qr_buf.getvalue()
        else:
            qr_bytes = b""

        # Also generate QR via free API as fallback image
        qr_api_url = f"https://api.qrserver.com/v1/create-qr-code/?size=300x300&data={urllib.parse.quote(link)}&color={color_hex.lstrip('#')}&bgcolor=ffffff&qzone=2&format=png"

        prog.progress(100, "Done!"); time.sleep(0.3); prog.empty()

        # Success banner
        st.markdown(f"""
        <div class="mt-qr-card">
            <div style="font-size:2rem;margin-bottom:8px;">🎉</div>
            <div style="font-family:'Syne',sans-serif;font-size:1.3rem;font-weight:800;color:#fff;margin-bottom:6px;">
                File Uploaded via {backend}!
            </div>
            <div style="color:rgba(255,255,255,0.5);font-size:0.85rem;margin-bottom:16px;">
                Share the QR code or copy the link. Anyone can scan to download instantly.
            </div>
            <span class="mt-qr-link">{link}</span>
            <div style="display:flex;gap:8px;flex-wrap:wrap;justify-content:center;margin-bottom:16px;">
                <span class="mt-info-chip">⏳ Expires: {expiry}</span>
                <span class="mt-info-chip">📁 {up.name}</span>
                <span class="mt-info-chip">⚖️ {size_mb:.2f} MB</span>
                <span class="mt-info-chip">☁️ {backend}</span>
            </div>
        </div>
        """, unsafe_allow_html=True)

        qr_col, info_col = st.columns([1, 1])
        with qr_col:
            if qr_bytes:
                st.image(qr_bytes, caption="📱 Scan to download file", use_container_width=True)
                st.download_button("⬇️ Save QR Code (PNG)", qr_bytes, "file_qr.png",
                                   "image/png", use_container_width=True, key="mt_qr_dl_qr")
            else:
                # Fallback: embed QR from free API
                st.markdown(f'<img src="{qr_api_url}" style="width:100%;border-radius:12px;" />', unsafe_allow_html=True)
                st.markdown(f"[⬇️ Download QR Code]({qr_api_url})", unsafe_allow_html=True)
                st.warning("Install qrcode[pil] for local QR generation.")

        with info_col:
            st.markdown(f"""
            <div style="padding:20px;background:rgba(10,14,30,0.8);border:1px solid rgba(255,255,255,0.07);
                border-radius:16px;">
                <div style="font-family:'JetBrains Mono',monospace;font-size:0.65rem;letter-spacing:3px;
                    color:#818cf8;text-transform:uppercase;margin-bottom:16px;">How to share</div>
                <div style="color:rgba(255,255,255,0.65);font-size:0.88rem;line-height:2.0;">
                    1️⃣ Copy the link above<br>
                    2️⃣ Or show the QR code<br>
                    3️⃣ Recipient scans → file downloads<br>
                    4️⃣ No account needed for either side<br>
                    <br>
                    <span style="color:rgba(255,255,255,0.35);font-size:0.78rem;">
                        🔒 Encrypted in transit · Free · No sign-up
                    </span>
                </div>
            </div>
            """, unsafe_allow_html=True)

        st.code(link, language=None)
        st.caption(f"✅ Uploaded via {backend} · Link expires: {expiry} · Copy and share freely")

# ═══════════════════════════════════════════════════════════
# MAIN ENTRY POINT
# ═══════════════════════════════════════════════════════════

def render_media_tools():
    """Main render function — called from app.py."""
    st.markdown(_MT_CSS, unsafe_allow_html=True)

    # Hero
    st.markdown("""
    <div class="mt-hero">
        <div class="mt-hero-title">🎨 Media Tools Suite</div>
        <div class="mt-hero-sub">
            AI Photo Editor · Image→PDF · PDF Converter · File QR Share
            — four powerful tools in one place
        </div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("← Back", key="mt_back"):
        st.session_state.app_mode = "chat"; st.rerun()

    tab1, tab2, tab3, tab4 = st.tabs([
        "🎨 AI Photo Editor",
        "🖼️→📄 Image → PDF",
        "📄→🔀 PDF Converter",
        "🔗 File QR Share",
    ])
    with tab1: _render_photo_editor()
    with tab2: _render_image_to_pdf()
    with tab3: _render_pdf_converter()
    with tab4: _render_file_qr_share()


# ── FREE API ADDITIONS ───────────────────────────────────────────────────────

def get_random_image_url(width: int = 800, height: int = 600, seed: int = None) -> str:
    """Get a free placeholder image URL from Picsum Photos (free, no key)."""
    if seed:
        return f"https://picsum.photos/seed/{seed}/{width}/{height}"
    return f"https://picsum.photos/{width}/{height}"


def get_color_name(hex_color: str) -> dict:
    """Get the name of a color from its hex code (TheColorAPI, free, no key)."""
    import urllib.request, json
    try:
        clean = hex_color.lstrip("#")
        req = urllib.request.Request(
            f"https://www.thecolorapi.com/id?hex={clean}",
            headers={"User-Agent": "ExamHelp/1.0"}
        )
        with urllib.request.urlopen(req, timeout=5) as r:
            data = json.loads(r.read().decode())
        return {"name": data.get("name", {}).get("value", ""),
                "hex": data.get("hex", {}).get("value", ""),
                "rgb": data.get("rgb", {}).get("value", ""),
                "hsl": data.get("hsl", {}).get("value", "")}
    except Exception:
        return {"name": hex_color, "hex": hex_color}


def generate_qr_url(data: str, size: int = 200) -> str:
    """Generate a QR code image URL using QR Server API (free, no key)."""
    import urllib.parse
    return f"https://api.qrserver.com/v1/create-qr-code/?size={size}x{size}&data={urllib.parse.quote(data)}"


def get_image_color_palette(image_url: str) -> list:
    """Extract dominant colors from an online image using ColorTag API (free)."""
    import urllib.request, urllib.parse, json
    try:
        url = f"https://colortagit.com/api/extract?img={urllib.parse.quote(image_url)}"
        req = urllib.request.Request(url, headers={"User-Agent": "ExamHelp/1.0"})
        with urllib.request.urlopen(req, timeout=6) as r:
            data = json.loads(r.read().decode())
        return data.get("colors", [])[:5]
    except Exception:
        return []


def shorten_url(long_url: str) -> str:
    """Shorten a URL using is.gd API (free, no key, for file share links)."""
    import urllib.request, urllib.parse
    try:
        api_url = f"https://is.gd/create.php?format=simple&url={urllib.parse.quote(long_url)}"
        req = urllib.request.Request(api_url, headers={"User-Agent": "ExamHelp/1.0"})
        with urllib.request.urlopen(req, timeout=5) as r:
            return r.read().decode().strip()
    except Exception:
        return long_url
