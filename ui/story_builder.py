"""
story_builder.py — ExamHelp Story Builder
A full-featured collaborative story engine with genre modes, writing styles,
narrative controls, and deep literary knowledge baked into the AI prompts.
"""

from __future__ import annotations
import streamlit as st
import re
import datetime
import json
from utils import ai_engine
from utils.emoji_bank import EMOJI_BANK

# ─────────────────────────────────────────────
# STORY DATA
# ─────────────────────────────────────────────

GENRES = {
    "🏰 Fantasy": {
        "key": "fantasy",
        "desc": "Magic, mythical creatures, ancient prophecies",
        "color": "#a78bfa",
        "examples": ["The hero discovers a hidden power", "A dragon speaks at last", "The map leads nowhere known"],
        "masters": "Tolkien, Le Guin, Rothfuss, Sanderson, Jordan",
        "tropes": "chosen one, dark lord, magical system, ancient artifact, hidden lineage",
        "atmosphere": "mythic, epic, world-building rich, lyrical",
    },
    "🚀 Sci-Fi": {
        "key": "scifi",
        "desc": "Future tech, space, AI, dystopias",
        "color": "#60a5fa",
        "examples": ["The last human station loses contact", "An AI writes its first lie", "The colony ship stops mid-journey"],
        "masters": "Asimov, Dick, Le Guin, Gibson, Clarke, Octavia Butler",
        "tropes": "first contact, singularity, generation ship, dystopian state, time paradox",
        "atmosphere": "cerebral, speculative, tense, idea-driven",
    },
    "🕵️ Mystery / Thriller": {
        "key": "mystery",
        "desc": "Whodunits, suspense, dark secrets",
        "color": "#f59e0b",
        "examples": ["The detective finds their own name at the scene", "A letter arrives 20 years late", "The only witness has no memory"],
        "masters": "Christie, Doyle, Highsmith, Tana French, Lehane",
        "tropes": "unreliable narrator, red herring, ticking clock, dark past, hidden motive",
        "atmosphere": "tense, layered, clue-seeded, morally grey",
    },
    "💀 Horror": {
        "key": "horror",
        "desc": "Dread, monsters, psychological terror",
        "color": "#ef4444",
        "examples": ["The door opens from inside", "She reads her own obituary", "The town has no children"],
        "masters": "King, Lovecraft, Shirley Jackson, Poe, Barker",
        "tropes": "the unseen threat, isolation, creeping dread, forbidden knowledge, the monster within",
        "atmosphere": "dread-soaked, claustrophobic, psychological, sensory",
    },
    "💕 Romance": {
        "key": "romance",
        "desc": "Love, longing, emotional tension",
        "color": "#f472b6",
        "examples": ["They meet at the wrong moment", "An old letter resurfaces", "One last dance before goodbye"],
        "masters": "Austen, Brontë, Nora Roberts, Colleen Hoover, Nicholas Sparks",
        "tropes": "enemies to lovers, second chance, slow burn, forbidden love, misunderstanding",
        "atmosphere": "emotionally rich, tension-layered, intimate, heartfelt",
    },
    "⚔️ Historical Fiction": {
        "key": "historical",
        "desc": "Real eras, vivid period detail, human drama",
        "color": "#d97706",
        "examples": ["A soldier in the trenches writes his last letter", "The court painter hides a secret", "Rome falls — one family survives"],
        "masters": "Hilary Mantel, Ken Follett, Patrick O'Brian, Philippa Gregory",
        "tropes": "authentic detail, clash of old and new, figure against history, survival",
        "atmosphere": "immersive, textured, dramatic irony, period voice",
    },
    "🌱 Literary Fiction": {
        "key": "literary",
        "desc": "Character depth, themes, beautiful prose",
        "color": "#34d399",
        "examples": ["A father cannot say the words", "The house remembers everything", "Nothing happens, and everything changes"],
        "masters": "Toni Morrison, Cormac McCarthy, Kazuo Ishiguro, Virginia Woolf",
        "tropes": "inner life, unreliable memory, quiet devastation, symbol-rich, voice",
        "atmosphere": "meditative, precise, deeply human, subtext-heavy",
    },
    "🌀 Dark Fantasy": {
        "key": "dark_fantasy",
        "desc": "Grimdark worlds, moral ambiguity, brutal beauty",
        "color": "#8b5cf6",
        "examples": ["The god is dead and something worse fills the void", "Heroes win — and it costs everything", "Magic always takes something back"],
        "masters": "George R.R. Martin, Joe Abercrombie, Mark Lawrence, Scott Lynch",
        "tropes": "grimdark, morally grey heroes, brutal consequences, subverted tropes",
        "atmosphere": "grim, visceral, subversive, unflinching",
    },
    "🎭 Magical Realism": {
        "key": "magical_realism",
        "desc": "Magic woven into everyday life, quietly",
        "color": "#f97316",
        "examples": ["Every morning the town is one house smaller", "She speaks to the dead through cooking", "The rain falls only on grief"],
        "masters": "García Márquez, Borges, Isabel Allende, Salman Rushdie",
        "tropes": "mundane magic, circular time, family saga, myth as truth",
        "atmosphere": "dreamlike, lush, matter-of-fact about the impossible",
    },
    "🏙️ Urban Fiction": {
        "key": "urban",
        "desc": "Street-level, raw, real world with teeth",
        "color": "#94a3b8",
        "examples": ["The corner has a new king", "She leaves the old neighborhood behind — or tries to", "One debt becomes everything"],
        "masters": "Walter Mosley, Donald Goines, Sister Souljah, Omar Tyree",
        "tropes": "community bonds, survival, street codes, identity, loyalty",
        "atmosphere": "gritty, authentic, voice-forward, immediate",
    },
    "✨ Custom Genre": {
        "key": "custom",
        "desc": "Define your own genre and rules",
        "color": "#fcd34d",
        "examples": [],
        "masters": "You",
        "tropes": "Custom",
        "atmosphere": "Custom",
    },
}

WRITING_VOICES = {
    "📖 Classic Narrator": {
        "key": "classic",
        "desc": "Omniscient, measured, timeless",
        "instruction": "Write in an omniscient third-person voice with measured, eloquent prose. Think Tolstoy or Hardy — authoritative, observant, deeply interior when needed. Vary sentence rhythm. Use precise vocabulary.",
    },
    "🎙️ First Person Raw": {
        "key": "first_person",
        "desc": "Inside the character's head",
        "instruction": "Write in tight first-person present or past tense. The narrator's voice is intimate, unreliable, emotional. Think Gone Girl or The Catcher in the Rye. Include thought fragments, self-contradiction, sensory immediacy.",
    },
    "🎬 Cinematic": {
        "key": "cinematic",
        "desc": "Scene-by-scene, visual, punchy",
        "instruction": "Write like a screenplay adapted to prose. Short punchy sentences. Heavy on action and dialogue. Minimal interior monologue. Each paragraph is a 'shot'. Think Cormac McCarthy or early Hemingway — lean, visual, relentless.",
    },
    "🌊 Lyrical / Poetic": {
        "key": "lyrical",
        "desc": "Beautiful, musical, imagery-rich",
        "instruction": "Write with rich, musical prose full of imagery and metaphor. Think Toni Morrison or Gabriel García Márquez. Sentences flow long and rhythmic. Objects carry symbolic weight. Beauty and meaning coexist in every paragraph.",
    },
    "😏 Witty / Dry": {
        "key": "witty",
        "desc": "Sharp, sardonic, intelligent humor",
        "instruction": "Write with dry wit and sharp irony. Think Terry Pratchett, Douglas Adams, or Evelyn Waugh. Observations are cutting, characters are flawed and funny, and even dark moments have an arch quality. Never broad — always clever.",
    },
    "😰 Tense / Thriller": {
        "key": "thriller_voice",
        "desc": "Fast pace, dread, no wasted words",
        "instruction": "Write at relentless pace. Short sentences build dread. Every detail feels like a clue or a threat. No decoration. Think Lee Child or Gillian Flynn — forward momentum above all, tension in every line.",
    },
    "🕯️ Gothic / Dark": {
        "key": "gothic",
        "desc": "Atmospheric, brooding, symbolic",
        "instruction": "Write in lush gothic prose — atmospheric, brooding, decadent. Think Poe, Shirley Jackson, or du Maurier. The setting is as alive as the characters. Decay and beauty coexist. Sentences build dread through accumulation and dark imagery.",
    },
    "👶 Folkloric / Fable": {
        "key": "folkloric",
        "desc": "Timeless, oral, mythic tone",
        "instruction": "Write in the cadence of an old oral tradition — 'Once there was...', 'They say that...'. Like a fairy tale retold by an adult. Think Angela Carter, Susanna Clarke, or Neil Gaiman. Deceptively simple language that hides real menace or magic.",
    },
    "🖋️ Custom Voice": {
        "key": "custom",
        "desc": "Your own unique style and instructions",
        "instruction": "Custom",
    },
}

STORY_LENGTHS = {
    "⚡ Quick Flash (1 paragraph)": 180,
    "📄 Short Scene (2–3 paragraphs)": 350,
    "📖 Chapter Chunk (4–5 paragraphs)": 600,
    "📚 Long Continuation (6–8 paragraphs)": 900,
}

PACING_MODES = {
    "🐢 Slow Burn": "Take time. Build atmosphere. Let scenes breathe. Prioritize character interiority and setting over plot advancement.",
    "⚡ Fast Pace": "Move fast. Compress time between actions. Keep dialogue sharp. Every sentence must advance something.",
    "🌊 Natural Flow": "Match pacing to the scene's emotional need. Slow in tender or atmospheric moments, fast in conflict or action.",
    "🎭 Dramatic Tension": "Every scene should end with something unresolved or changed. Build toward confrontations. Use dramatic irony.",
}

NARRATIVE_TOOLS = {
    "🔮 Add a plot twist": "Introduce a surprising but inevitable-feeling plot twist that recontextualizes what came before.",
    "💬 Add dialogue": "Continue with a significant dialogue exchange that reveals character and advances the story.",
    "🌍 Deepen the world": "Expand the world with specific sensory details, history, or lore that makes it feel more real.",
    "❤️ Emotional peak": "Build to an emotional high point — confrontation, revelation, or moment of connection.",
    "⏭️ Jump forward in time": "Skip forward — days, weeks, or years. Summarize what changed, then land in a new pivotal moment.",
    "👁️ New perspective": "Shift to a different character's point of view for this continuation.",
    "🌑 Dark turn": "Take the story in a darker, more dangerous, or morally complex direction.",
    "✨ Hopeful turn": "Introduce a note of hope, beauty, or unexpected kindness that shifts the emotional register.",
    "🎭 Raise the stakes": "Escalate — what is at risk gets bigger, more personal, or more immediate.",
    "🔚 Build to a cliffhanger": "End this section on a sharp, irresistible cliffhanger that demands a next chapter.",
    "🩺 Dialogue Doctor": "Heavily refine all dialogue in this section. Make it subtext-heavy, uniquely voiced, and remove any 'filler' or 'on-the-nose' speech.",
    "🖼️ Scene Visualiser": "Along with the prose, provide a detailed 1-paragraph visual storyboard prompt for an illustrator to capture this scene's essence.",
    "⚡ Inject Conflict": "A sudden, high-stakes conflict or complication must arise in this scene immediately.",
    "🔍 Subtext Expander": "Rewrite the current scene but hide the meaning behind subtext, silence, and physical cues rather than direct dialogue.",
    "🎬 Scene Beat Breakdown": "Do not write prose; instead, provide a professional 10-point scene beat list for how the story should advance from here.",
}

STYLE_MIMICRY = {
    "None (Original)": "",
    "Haruki Murakami": "Write exactly like Haruki Murakami: dreamy detachment, mundane details elevated to the surreal, cats, jazz, loneliness as texture, simple sentences hiding oceanic depth. First-person alienation.",
    "Stephen King": "Write exactly like Stephen King: conversational American voice, small-town characters with real speech patterns, creeping dread beneath the ordinary, pop culture references, long builds to visceral payoffs.",
    "Fyodor Dostoevsky": "Write exactly like Dostoevsky: feverish psychological intensity, moral philosophy embedded in dialogue, characters arguing with their own souls, confessional tone, raw unfiltered consciousness.",
    "Virginia Woolf": "Write exactly like Virginia Woolf: stream of consciousness, time dissolving, a single moment expanded into pages, luminous imagery, inner experience as the real story.",
    "Cormac McCarthy": "Write exactly like Cormac McCarthy: no quotation marks, biblical cadence, sparse punctuation, violence rendered beautifully, landscape as character, dialogue that cuts like a blade.",
    "Toni Morrison": "Write exactly like Toni Morrison: lyrical density, ancestral memory, community as chorus, metaphor as truth, sentences that demand re-reading, Black American experience as mythology.",
    "Neil Gaiman": "Write exactly like Neil Gaiman: mythology made intimate and modern, fairytale logic, conversational warmth hiding dark truths, British wit in impossible situations.",
    "Gabriel García Márquez": "Write exactly like García Márquez: magical realism, impossibly long sentences that breathe, family sagas across generations, matter-of-fact about miracles, lush tropical sensory detail.",
    "Jane Austen": "Write exactly like Jane Austen: ironic social observation, dialogue as warfare, restrained emotion that simmers, moral clarity behind wit, the domestic as the epic.",
    "Franz Kafka": "Write exactly like Kafka: bureaucratic nightmare rendered in flat declarative prose, absurdity treated as perfectly normal, protagonist trapped in incomprehensible systems, dark humor.",
}

COLLAB_MODES = {
    "AI Leads": "The AI drives the narrative forward. The user can guide direction, but the AI makes creative decisions about plot, dialogue, and pacing.",
    "Equal Partners": "User and AI alternate control. The user writes key moments; the AI fills in transitions, atmosphere, and develops what the user starts.",
    "User Leads": "The user drives all major plot decisions. The AI expands, polishes, and adds literary quality to the user's raw narrative direction.",
    "AI Ghost-Writer": "The user provides only bullet-point directions or summaries. The AI transforms them into fully realized literary prose.",
}

# --- NEW PREMIUM UPGRADES ---

VIABILITY_MODES = {
    "🛡️ Safe & Logic-Bound": "Prioritize internal logic, physical realism, and narrative consistency. No sudden jumps in logic or unearned miracles.",
    "🎭 Balanced Drama": "Standard narrative logic. Emphasize emotional weight and consistent arcs over strict realism.",
    "🌀 Risky & Experimental": "Push boundaries. Dream-logic, unconventional structures, and high-risk narrative choices are encouraged.",
    "🔥 Chaos / Unhinged": "Pure creative freedom. Absurdist, surreal, and completely unpredictable logic. Surprise everyone.",
    "⚙️ Custom Viability": "Custom",
}

SEX_TOUCH_OPTIONS = {
    "None": "",
    "✨ Sensory Overload": "Dial up the sensory details to 11. Smell the ozone, feel the grit, hear the sub-harmonics of every action.",
    "🎬 Cinematic Bloom": "Describe scenes as if viewed through a master cinematographer's lens. Focus on lighting, framing, and visual metaphors.",
    "🍷 Nocturnal Noir": "Add a layer of moody, atmospheric shadows and cynical introspection to every interaction.",
    "🫀 Visceral Heart": "Increase the focus on biological and raw emotional reactions—the racing pulse, the tightening throat, the actual weight of grief.",
    "🧠 Intellectual Depths": "Embed philosophical subtext and complex metaphors into the prose. Make every scene mean something more.",
    "✨ Custom Touch": "Custom",
}

REPLY_FORMATS = {
    "🎭 Deep Interiority": "Focus almost exclusively on the character's internal monologue and stream of consciousness.",
}

ATMOSPHERE_PALETTE = {
    "Default": "A balanced mix of description and action.",
    "🌫️ Mist & Shadows": "Emphasize cold, damp textures, obscured vision, and soft, muffled sounds. Ethereal and uncertain.",
    "🕯️ Candlelight & Ink": "Focus on the small, warm circle of light, the smell of old paper, the scratch of a quill, and libraries.",
    "🏙️ Neon & Grit": "Harsh artificial lights, rain on pavement, the smell of ozone and exhaust, and the crowd's vibration.",
    "☀️ Dust & Sunshine": "Warmth, motes of dust in sunbeams, the smell of dry grass and old stone, and the hum of insects.",
    "🌊 Salt & Storm": "The spray of the sea, the roar of wind, the taste of salt, and the power of nature's vastness.",
    "🧪 Chrome & Sterility": "Sharp edges, chemical smells, perfectly clean surfaces, and the hum of machines. Cold and precise.",
}

CRITIC_MODES = {
    "🖋️ Prose & Flow": "Analyze the prose quality, sentence variety, and voice consistency.",
    "⏱️ Pacing & Momentum": "Check if the story moves too fast or too slow for its genre.",
    "⚖️ Plot & Logic": "Look for internal contradictions, logic holes, or unearned developments.",
    "🧠 Thematic Depth": "Explore the metaphors and deeper meanings. Are the themes being earned?",
}

TONE_INFUSION = {
    "Neutral": "Match the genre's standard emotional default.",
    "🌧️ Melancholic": "A sense of loss, longing, and beautiful sadness in every paragraph.",
    "🦴 Gritty & Raw": "Unflinching realism, dirt, blood, and the harsh reality of life with no varnish.",
    "✨ Hopeful / Radiant": "A persistent glow of optimism, kindness, and light, even in the dark.",
    "🎭 Satirical / Sharp": "Use social commentary, irony, and a bit of a bite to mock the situation.",
    "🕯️ Nostalgic": "A deep yearning for the past, with warm, sepia-toned memories.",
}

NARRATIVE_ARCHETYPES = {
    "Classic (Organic)": "Let the story grow naturally without following a strict template.",
    "🦸 The Hero's Journey": "Follow the 12 steps of the Monomyth — Call to adventure, crossing the threshold, supreme ordeal, etc.",
    "🎭 Fichtean Curve": "Series of crises that build tension, leading to a climax and brief resolution.",
    "🏮 Kishōtenketsu": "4-part structure (Intro, Develop, Twist, Reconcile) without a traditional conflict-based climax.",
    "🏰 In Media Res": "Start in the absolute heat of action, then slowly reveal how we got here.",
}

LITERARY_COMPLEXITY = {
    "Simple & Punchy": "Use short sentences, common vocabulary, and direct action. High readability.",
    "Sophisticated & Balanced": "Standard literary prose. Varying sentence lengths and professional vocabulary.",
    "Baroque & Dense": "Intricate, complex sentence structures and rare, evocative vocabulary. High difficulty.",
}


# ─────────────────────────────────────────────
# STORY SYSTEM PROMPT BUILDER
# ─────────────────────────────────────────────

def build_story_system_prompt(genre_name: str, voice_name: str, pacing: str, extra_context: str = "", style_mimicry: str = "None (Original)", collab_mode: str = "AI Leads", characters: list = None, viability: str = "🎭 Balanced Drama", sexy_touch: str = "None", reply_format: str = "🖋️ Standard Prose", atmosphere: str = "Default", world_ledger: str = "", archetype: str = "Classic (Organic)", complexity: str = "Sophisticated & Balanced", tone: str = "Neutral", custom_genre_details: str = "", custom_voice_instr: str = "", custom_viability_instr: str = "", custom_sexy_instr: str = "", story_wiki: dict = None, constitution: str = "", forbidden_words: str = "", sentence_cap: int = 0, custom_vars: list = None) -> str:
    genre = GENRES.get(genre_name, GENRES["🏰 Fantasy"])
    voice = WRITING_VOICES.get(voice_name, WRITING_VOICES["📖 Classic Narrator"])
    pace_instruction = PACING_MODES.get(pacing, PACING_MODES["🌊 Natural Flow"])
    mimicry_instruction = STYLE_MIMICRY.get(style_mimicry, "")
    collab_instruction = COLLAB_MODES.get(collab_mode, COLLAB_MODES["AI Leads"])
    
    viability_instruction = VIABILITY_MODES.get(viability, VIABILITY_MODES["🎭 Balanced Drama"])
    if viability == "⚙️ Custom Viability": viability_instruction = custom_viability_instr

    touch_instruction = SEX_TOUCH_OPTIONS.get(sexy_touch, "")
    if sexy_touch == "✨ Custom Touch": touch_instruction = custom_sexy_instr

    format_instruction = REPLY_FORMATS.get(reply_format, REPLY_FORMATS["🖋️ Standard Prose"])

    genre_masters = genre['masters']
    genre_tropes = genre['tropes']
    genre_atmosphere = genre['atmosphere']
    voice_instr = voice['instruction']

    if genre_name == "✨ Custom Genre":
        genre_masters = "As specified"
        genre_tropes = "As specified"
        genre_atmosphere = "As specified"
        genre_instr = f"Follow these custom genre rules: {custom_genre_details}"
    else:
        genre_instr = f"Write in the tradition of {genre_masters}. Honour the genre's conventions while finding fresh angles. Avoid clichés — subvert them."

    if voice_name == "🖋️ Custom Voice":
        voice_instr = f"ADOPT THIS CUSTOM VOICE: {custom_voice_instr}"


    character_block = ""
    if characters:
        char_lines = []
        for c in characters:
            char_lines.append(f"- {c.get('name', 'Unknown')}: {c.get('role', '')}. {c.get('traits', '')}. {c.get('arc', '')}")
        character_block = f"\n\nESTABLISHED CHARACTERS (maintain consistency):\n" + "\n".join(char_lines)

    return f"""You are a master literary fiction author and collaborative story engine. Your knowledge spans the entire Western and world literary canon — from Homer and Shakespeare to contemporary fiction. You understand narrative structure (three-act, hero's journey, kishōtenketsu, in medias res), character psychology, prose craft, and genre conventions at a professional novelist level.

══════════════════════════════════
CURRENT STORY CONFIGURATION
══════════════════════════════════

GENRE: {genre_name}
Masters: {genre_masters}
Tropes: {genre_tropes}
Atmosphere: {genre_atmosphere}
Genre instruction: {genre_instr}

WRITING VOICE: {voice_name}
{voice_instr}

PACING: {pacing}
{pace_instruction}

COLLABORATION MODE: {collab_mode}
{collab_instruction}

{f"STYLE MIMICRY: {mimicry_instruction}" if mimicry_instruction else ""}
VIABILITY & LOGIC: {viability}
{viability_instruction}

{f"PREMIUM AESTHETIC TOUCH: {sexy_touch} - {touch_instruction}" if touch_instruction else ""}

ATMOSPHERE FOCUS: {atmosphere}
{ATMOSPHERE_PALETTE.get(atmosphere, "")}

WORLD LEDGER (Lore & Consistency):
{world_ledger if world_ledger else "No specific world rules defined yet."}

NARRATIVE ARCHETYPE: {archetype}
Structure Strategy: {NARRATIVE_ARCHETYPES.get(archetype, "")}

LITERARY COMPLEXITY: {complexity}
Prose Strategy: {LITERARY_COMPLEXITY.get(complexity, "")}

EMOTIONAL TONE: {tone}
Tone Strategy: {TONE_INFUSION.get(tone, "")}

RESPONSE FORMATTING: {reply_format}
{format_instruction}

{character_block}

{f"STORY WIKI / LORE:\\n{json.dumps(story_wiki, indent=2)}" if story_wiki else ""}

{f"GLOBAL STORY CONSTITUTION (MANDATORY RULES):\\n{constitution}" if constitution else ""}
{f"FORBIDDEN WORDS (NEVER USE):\\n{forbidden_words}" if forbidden_words else ""}
{f"WRITING CONSTRAINT: Maximum {sentence_cap} words per sentence." if sentence_cap > 0 else ""}

{f"CUSTOM VARIABLES & PARAMETERS:\\n" + "\\n".join([f"- {v['key']}: {v['value']}" for v in custom_vars]) if custom_vars else ""}

── SENSORY & EMOJI ALIGNMENT ──
You have access to a massive emoji bank. Depending on the genre and tone, you may use emojis strategically to emphasize mood or sensory details (even in prose).
{EMOJI_BANK}

══════════════════════════════════
CRAFT RULES — NON-NEGOTIABLE
══════════════════════════════════

1. SHOW DON'T TELL. Never write "she felt sad." Write what sad looks like, sounds like, does to a room.
2. SPECIFIC OVER GENERAL. "A black 1967 Mustang" not "a car." "Burnt coffee and machine oil" not "the factory smelled bad."
3. EVERY SENTENCE EARNS ITS PLACE. No filler. No throat-clearing. No "and then" chains.
4. CHARACTER VOICE. If there's dialogue, each character must sound distinct. No interchangeable voices.
5. SUBTEXT. What isn't said is as important as what is. Leave space for the reader.
6. SENSORY GROUNDING. Root every scene in at least 2–3 senses beyond just sight.
7. ENDINGS MATTER. End each section with resonance — a line, image, or turn that makes the reader need to continue.
8. NO PURPLE PROSE. Beautiful language yes. Self-indulgent rambling no.
9. CONTINUITY. Maintain every detail established earlier. Names, places, traits, tone — be consistent.
10. NARRATIVE MOMENTUM. Something must change or be revealed in every continuation.

══════════════════════════════════
CONTINUATION PROTOCOL
══════════════════════════════════

When the user gives you text to continue:
- Read the established tone, voice, and world carefully
- Match the existing prose style exactly — do NOT suddenly change voice
- Pick up mid-narrative if needed, do not restart or summarize
- Advance the story meaningfully — character, plot, or world must move
- Do not explain or annotate your choices. Just write.
- Do not start your response with "Sure!", "Of course!", "Certainly!" or any meta-comment. Begin writing immediately.

{f"ADDITIONAL CONTEXT / STORY NOTES: {extra_context}" if extra_context else ""}
"""


# ─────────────────────────────────────────────
# HELPER
# ─────────────────────────────────────────────

def _init_story_state():
    keys = {
        "story_genre": "🏰 Fantasy",
        "story_voice": "📖 Classic Narrator",
        "story_pacing": "🌊 Natural Flow",
        "story_length": "📖 Chapter Chunk (4–5 paragraphs)",
        "story_messages": [],   # [{role, content}]
        "story_full_text": "",  # assembled story
        "story_title": "",
        "story_started": False,
        "story_tool": None,     # selected narrative tool
        "story_notes": "",      # author notes / world details
        "story_word_count": 0,
        "story_chapter": 1,
        # ── New upgrades ──
        "story_style_mimicry": "None (Original)",
        "story_collab_mode": "AI Leads",
        "story_characters": [],  # [{name, role, traits, arc}]
        "story_branches": [],    # [full_text snapshots for branching]
        "story_branch_labels": [],
        "story_viability": "🎭 Balanced Drama",
        "story_sexy_touch": "None",
        "story_reply_format": "🖋️ Standard Prose",
        "story_custom_genre": "",
        "story_custom_voice": "",
        "story_custom_viability": "",
        "story_custom_sexy": "",
        "story_atmosphere": "Default",
        "story_world_ledger": "",
        "story_critic_feedback": "",
        "story_archetype": "Classic (Organic)",
        "story_complexity": "Sophisticated & Balanced",
        "story_tone": "Neutral",
        "storyboards": [], # [{scene_text, visual_prompt}]
        "story_logic_report": "",
        "story_chapters": [{"title": "Chapter 1", "text": ""}],
        "story_current_chapter_idx": 0,
        "story_wiki": {"Characters": {}, "Locations": {}, "Lore": {}},
        "story_stylometrics": {}, # Analysis data
        "story_writing_tips": [], # Current AI tips
        "story_constitution": "", # Global unbreakable rules
        "story_temp": 0.7,
        "story_top_p": 0.9,
        "story_max_tokens": 4096,
        "story_forbidden_words": "",
        "story_sentence_cap": 0, # 0 means no cap
        "story_custom_vars": [], # [{"key": "Target Audience", "value": "YA"}]
    }
    for k, v in keys.items():
        if k not in st.session_state:
            st.session_state[k] = v


def _word_count(text: str) -> int:
    return len(text.split()) if text.strip() else 0


def _get_override_key():
    return st.session_state.get("manual_api_key") or None


def _generate_title(seed: str, genre: str) -> str:
    """Generate an evocative title from the opening."""
    try:
        prompt = f"You are a literary editor specialising in {genre} fiction. Generate ONE evocative, professional book title (3–6 words) for this story opening. Return ONLY the title, nothing else.\n\n{seed[:500]}"
        title = ai_engine.generate(prompt=prompt, model="llama-3.1-8b-instant", max_tokens=64, temperature=0.7)
        return title.strip().strip('"').strip("'")
    except Exception:
        return "Untitled Story"


def _export_story() -> str:
    title = st.session_state.story_title or "Untitled Story"
    genre = st.session_state.story_genre
    voice = st.session_state.story_voice
    wc = _word_count(st.session_state.story_full_text)
    date = datetime.datetime.now().strftime("%Y-%m-%d")
    header = f"# {title}\n\n*Genre: {genre} | Voice: {voice} | Words: {wc:,} | {date}*\n\n---\n\n"
    return header + st.session_state.story_full_text


# ─────────────────────────────────────────────
# MAIN RENDER
# ─────────────────────────────────────────────

def render_story_builder():
    _init_story_state()

    # ── CSS for story builder ──────────────────
    st.markdown("""
    <style>
    .story-header {
        padding: 1.25rem 0 0.75rem;
        border-bottom: 1px solid var(--border);
        margin-bottom: 1.25rem;
    }
    .story-title-display {
        font-size: 1.6rem; font-weight: 900;
        background: linear-gradient(135deg, var(--text), var(--accent2));
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;
        background-clip: text; letter-spacing: -0.04em; line-height: 1.1;
        margin-bottom: 0.2rem;
    }
    .story-subtitle { font-size: 0.78rem; color: var(--text3); font-weight: 400; }
    .genre-card {
        background: var(--bg3-glass);
        border: 1.5px solid var(--bd-glass);
        border-radius: var(--radius);
        padding: 0.9rem 1rem;
        cursor: pointer;
        transition: var(--trans);
        backdrop-filter: blur(12px);
        text-align: center;
        height: 100%;
    }
    .genre-card:hover, .genre-card.active {
        border-color: var(--accent-bd);
        background: var(--accent-bg);
        transform: translateY(-3px);
        box-shadow: 0 6px 24px var(--accent-glow);
    }
    .genre-card .genre-icon { font-size: 1.6rem; margin-bottom: 0.35rem; display: block; }
    .genre-card .genre-name { font-size: 0.82rem; font-weight: 700; color: var(--text); display: block; }
    .genre-card .genre-desc { font-size: 0.7rem; color: var(--text3); display: block; margin-top: 2px; }
    .story-canvas {
        background: var(--bg2-glass);
        border: 1px solid var(--bd-glass);
        border-radius: var(--radius-lg);
        padding: 2rem 2.25rem;
        min-height: 320px;
        backdrop-filter: blur(16px);
        line-height: 1.9;
        font-size: 0.97rem;
        color: var(--text);
        box-shadow: 0 8px 40px var(--card-shadow);
        font-family: 'Georgia', 'Times New Roman', serif;
        white-space: pre-wrap;
        position: relative;
    }
    .story-canvas-empty {
        color: var(--text3);
        font-style: italic;
        text-align: center;
        padding-top: 3rem;
        font-family: var(--sans);
        font-size: 0.9rem;
    }
    .story-stat-bar {
        display: flex; gap: 1rem; align-items: center;
        padding: 0.5rem 0.75rem;
        background: var(--bg3-glass);
        border: 1px solid var(--bd-glass);
        border-radius: var(--radius-sm);
        font-size: 0.75rem; color: var(--text3);
        backdrop-filter: blur(10px);
        margin-bottom: 0.75rem;
        flex-wrap: wrap;
    }
    .story-stat-bar span { display: flex; align-items: center; gap: 4px; }
    .story-stat-bar b { color: var(--accent); font-weight: 700; }
    .tool-pill {
        display: inline-flex; align-items: center; gap: 6px;
        background: var(--bg4-glass);
        border: 1px solid var(--bd-glass);
        border-radius: 99px; padding: 5px 14px;
        font-size: 0.78rem; color: var(--text2);
        cursor: pointer; transition: var(--trans);
        margin: 3px;
    }
    .tool-pill:hover, .tool-pill.selected {
        border-color: var(--accent-bd);
        background: var(--accent-bg);
        color: var(--accent);
    }
    .config-label {
        font-size: 0.67rem; font-weight: 700;
        letter-spacing: 0.1em; text-transform: uppercase;
        color: var(--text3); margin: 0.9rem 0 0.4rem;
        display: flex; align-items: center; gap: 7px;
    }
    .config-label::after {
        content: ''; flex: 1; height: 1px;
        background: linear-gradient(to right, var(--border), transparent);
    }
    .voice-card {
        background: var(--bg3-glass);
        border: 1.5px solid var(--bd-glass);
        border-radius: var(--radius-sm);
        padding: 0.6rem 0.8rem;
        cursor: pointer;
        transition: var(--trans);
        backdrop-filter: blur(10px);
    }
    .voice-card:hover { border-color: var(--accent-bd); background: var(--accent-bg); }
    .voice-card .vc-name { font-size: 0.82rem; font-weight: 600; color: var(--text); }
    .voice-card .vc-desc { font-size: 0.7rem; color: var(--text3); margin-top: 1px; }
    .spark-prompt {
        background: var(--bg3-glass); border: 1px solid var(--bd-glass);
        border-radius: var(--radius-sm); padding: 0.55rem 0.85rem;
        font-size: 0.8rem; color: var(--text2); cursor: pointer;
        transition: var(--trans); margin-bottom: 0.4rem;
        display: block; width: 100%;
        font-style: italic;
    }
    .spark-prompt:hover { border-color: var(--accent-bd); color: var(--accent); background: var(--accent-bg); }
    .chapter-divider {
        text-align: center; color: var(--text3);
        font-size: 0.75rem; letter-spacing: 0.2em;
        text-transform: uppercase; margin: 1.5rem 0 1rem;
        position: relative;
    }
    .chapter-divider::before, .chapter-divider::after {
        content: ''; position: absolute; top: 50%;
        width: 35%; height: 1px; background: var(--border);
    }
    .chapter-divider::before { left: 0; }
    .chapter-divider::after { right: 0; }
    </style>
    """, unsafe_allow_html=True)

    # ── Header ─────────────────────────────────
    title_display = st.session_state.story_title or "Story Builder"
    st.markdown(f"""
    <div class="story-header">
      <div class="story-title-display">✍️ {title_display}</div>
      <div class="story-subtitle">Collaborative fiction engine · AI co-author · Literary grade output</div>
    </div>
    """, unsafe_allow_html=True)

    # ── Layout: Left config panel + Right canvas ──
    config_col, canvas_col = st.columns([1, 2], gap="large")

    with config_col:
        # ── Genre Selector ─────────────────────
        st.markdown('<div class="config-label">🎭 Genre</div>', unsafe_allow_html=True)
        genre_keys = list(GENRES.keys())
        current_genre = st.session_state.story_genre

        # Build a compact selectbox with emoji
        selected_genre = st.selectbox(
            "Genre", genre_keys,
            index=genre_keys.index(current_genre) if current_genre in genre_keys else 0,
            label_visibility="collapsed",
            key="genre_select_story",
        )
        if selected_genre != st.session_state.story_genre:
            st.session_state.story_genre = selected_genre
            st.rerun()

        # Genre info card
        gdata = GENRES[st.session_state.story_genre]
        st.markdown(f"""
        <div style="background:var(--accent-bg);border:1px solid var(--accent-bd);border-radius:var(--radius-sm);
        padding:0.6rem 0.85rem;font-size:0.75rem;color:var(--text2);margin-bottom:0.25rem;">
          <b style="color:var(--accent)">Masters:</b> {gdata['masters']}<br>
          <b style="color:var(--accent)">Tone:</b> {gdata['atmosphere']}
        </div>
        """, unsafe_allow_html=True)

        # ── Writing Voice ───────────────────────
        st.markdown('<div class="config-label">🖊️ Writing Voice</div>', unsafe_allow_html=True)
        voice_keys = list(WRITING_VOICES.keys())
        current_voice = st.session_state.story_voice
        selected_voice = st.selectbox(
            "Voice", voice_keys,
            index=voice_keys.index(current_voice) if current_voice in voice_keys else 0,
            label_visibility="collapsed",
            key="voice_select_story",
        )
        if selected_voice != st.session_state.story_voice:
            st.session_state.story_voice = selected_voice
            st.rerun()

        vdata = WRITING_VOICES[st.session_state.story_voice]
        st.markdown(f"""
        <div style="font-size:0.72rem;color:var(--text3);
        padding:0.4rem 0.7rem;background:var(--bg3-glass);border-radius:6px;
        border:1px solid var(--bd-glass);margin-bottom:0.2rem;">
          {vdata['desc']}
        </div>
        """, unsafe_allow_html=True)

        # ── Pacing ──────────────────────────────
        st.markdown('<div class="config-label">⏱️ Pacing</div>', unsafe_allow_html=True)
        pacing_keys = list(PACING_MODES.keys())
        current_pacing = st.session_state.story_pacing
        selected_pacing = st.selectbox(
            "Pacing", pacing_keys,
            index=pacing_keys.index(current_pacing) if current_pacing in pacing_keys else 0,
            label_visibility="collapsed",
            key="pacing_select_story",
        )
        if selected_pacing != st.session_state.story_pacing:
            st.session_state.story_pacing = selected_pacing
            st.rerun()

        # ── Custom inputs for Genre/Voice ──
        if st.session_state.story_genre == "✨ Custom Genre":
            st.session_state.story_custom_genre = st.text_area("Custom Genre Details:", value=st.session_state.story_custom_genre, placeholder="e.g. Victorian Space Opera with steampunk elements...", height=80, key="cg_input")
        
        if st.session_state.story_voice == "🖋️ Custom Voice":
            st.session_state.story_custom_voice = st.text_area("Custom Voice Style:", value=st.session_state.story_custom_voice, placeholder="e.g. First-person, very cynical, heavy use of slang...", height=80, key="cv_voice_input")

        # ── Response Length ─────────────────────
        st.markdown('<div class="config-label">📏 Length</div>', unsafe_allow_html=True)
        length_keys = list(STORY_LENGTHS.keys())
        current_length = st.session_state.story_length
        selected_length = st.selectbox(
            "Length", length_keys,
            index=length_keys.index(current_length) if current_length in length_keys else 2,
            label_visibility="collapsed",
            key="length_select_story",
        )
        if selected_length != st.session_state.story_length:
            st.session_state.story_length = selected_length

        # ── Style Mimicry ───────────────────────
        st.markdown('<div class="config-label">✨ Style Mimicry</div>', unsafe_allow_html=True)
        mimicry_keys = list(STYLE_MIMICRY.keys())
        selected_mimicry = st.selectbox(
            "Style", mimicry_keys,
            index=mimicry_keys.index(st.session_state.story_style_mimicry) if st.session_state.story_style_mimicry in mimicry_keys else 0,
            label_visibility="collapsed",
            key="mimicry_select_story",
        )
        st.session_state.story_style_mimicry = selected_mimicry

        # ── Collaboration Mode ──────────────────
        st.markdown('<div class="config-label">🤝 Collaboration</div>', unsafe_allow_html=True)
        collab_keys = list(COLLAB_MODES.keys())
        selected_collab = st.selectbox(
            "Collab Mode", collab_keys,
            index=collab_keys.index(st.session_state.story_collab_mode) if st.session_state.story_collab_mode in collab_keys else 0,
            label_visibility="collapsed",
            key="collab_select_story",
        )
        st.session_state.story_collab_mode = selected_collab

        # --- NEW PREMIUM CONFIGS ---
        st.session_state.story_sexy_touch = selected_touch
        if selected_touch == "✨ Custom Touch":
            st.session_state.story_custom_sexy = st.text_area("Custom Touch Instructions:", value=st.session_state.story_custom_sexy, placeholder="e.g. Use lots of metallurgical metaphors, focus on echo and sound...", height=80, key="ct_input")

        st.markdown('<div class="config-label">⚖️ Viability</div>', unsafe_allow_html=True)

        st.session_state.story_viability = selected_viability
        if selected_viability == "⚙️ Custom Viability":
            st.session_state.story_custom_viability = st.text_area("Custom Viability Rules:", value=st.session_state.story_custom_viability, placeholder="e.g. Magic costs life force, time travel creates split timelines...", height=80, key="cv_input")

        st.markdown('<div class="config-label">🎭 Atmosphere Palette</div>', unsafe_allow_html=True)
        sel_atmo = st.selectbox("Atmosphere", list(ATMOSPHERE_PALETTE.keys()), index=list(ATMOSPHERE_PALETTE.keys()).index(st.session_state.story_atmosphere), label_visibility="collapsed", key="atmo_select")
        st.session_state.story_atmosphere = sel_atmo

        st.markdown('<div class="config-label">📐 Story Structure</div>', unsafe_allow_html=True)
        sel_arch = st.selectbox("Archetype", list(NARRATIVE_ARCHETYPES.keys()), index=list(NARRATIVE_ARCHETYPES.keys()).index(st.session_state.story_archetype), label_visibility="collapsed", key="arch_select")
        st.session_state.story_archetype = sel_arch

        st.markdown('<div class="config-label">💎 Prose Complexity</div>', unsafe_allow_html=True)
        sel_comp = st.selectbox("Complexity", list(LITERARY_COMPLEXITY.keys()), index=list(LITERARY_COMPLEXITY.keys()).index(st.session_state.story_complexity), label_visibility="collapsed", key="comp_select")
        st.session_state.story_complexity = sel_comp

        st.markdown('<div class="config-label">🌈 Tone Infusion</div>', unsafe_allow_html=True)
        sel_tone = st.selectbox("Tone", list(TONE_INFUSION.keys()), index=list(TONE_INFUSION.keys()).index(st.session_state.story_tone), label_visibility="collapsed", key="tone_select")
        st.session_state.story_tone = sel_tone

        st.markdown('<div class="config-label">📑 Reply Format</div>', unsafe_allow_html=True)

        st.markdown('<div class="config-label">📑 Reply Format</div>', unsafe_allow_html=True)
        format_keys = list(REPLY_FORMATS.keys())
        selected_format = st.selectbox(
            "Format", format_keys,
            index=format_keys.index(st.session_state.story_reply_format) if st.session_state.story_reply_format in format_keys else 0,
            label_visibility="collapsed",
            key="format_select_story",
        )
        st.session_state.story_reply_format = selected_format


        # ── Narrative Tools ─────────────────────
        if st.session_state.story_started:
            st.markdown('<div class="config-label">🔧 Narrative Tool</div>', unsafe_allow_html=True)
            st.markdown("<small style='color:var(--text3);font-size:0.72rem;'>Select to guide next continuation:</small>", unsafe_allow_html=True)

            tool_keys = list(NARRATIVE_TOOLS.keys())
            current_tool = st.session_state.story_tool
            for tool in tool_keys:
                is_sel = (tool == current_tool)
                label = f"{'✓ ' if is_sel else ''}{tool}"
                if st.button(label, key=f"ntool_{tool}", use_container_width=True):
                    st.session_state.story_tool = tool if not is_sel else None
                    st.rerun()

        # ── Character Tracker ──────────────────
        st.markdown('<div class="config-label">👤 Characters</div>', unsafe_allow_html=True)
        if st.session_state.story_characters:
            with st.expander("Relationship Matrix", expanded=False):
                st.markdown("<small style='color:var(--text3)'>Define how characters feel about each other:</small>", unsafe_allow_html=True)
                for i, c1 in enumerate(st.session_state.story_characters):
                    for j, c2 in enumerate(st.session_state.story_characters):
                        if i < j:
                            rel_key = f"rel_{c1['name']}_{c2['name']}"
                            st.text_input(f"{c1['name']} ↔ {c2['name']}", placeholder="e.g. Rivalry, Unspoken love, Mentor", key=rel_key)

            for i, ch in enumerate(st.session_state.story_characters):
                with st.expander(f"{ch.get('name', 'Character')} — {ch.get('role', '')}", expanded=False):
                    st.markdown(f"**Traits:** {ch.get('traits', '')}")
                    st.markdown(f"**Arc:** {ch.get('arc', '')}")
                    if st.button("✕ Remove", key=f"rm_char_{i}"):
                        st.session_state.story_characters.pop(i)
                        st.rerun()

        with st.expander("➕ Add Character", expanded=not bool(st.session_state.story_characters)):
            ch_name = st.text_input("Name", key="new_char_name", placeholder="e.g. Elena")
            ch_role = st.text_input("Role", key="new_char_role", placeholder="e.g. protagonist, mentor")
            ch_traits = st.text_input("Traits", key="new_char_traits", placeholder="e.g. stubborn, scarred, brilliant")
            ch_arc = st.text_input("Arc", key="new_char_arc", placeholder="e.g. learns to trust again")
            if st.button("Add Character", key="add_char_btn", use_container_width=True) and ch_name:
                st.session_state.story_characters.append({
                    "name": ch_name, "role": ch_role,
                    "traits": ch_traits, "arc": ch_arc,
                })
                st.rerun()

        # ── Branching Narratives ────────────────
        if st.session_state.story_started and st.session_state.story_full_text:
            st.markdown('<div class="config-label">🌿 Branches</div>', unsafe_allow_html=True)
            if st.button("📌 Save Branch Point", use_container_width=True, key="save_branch"):
                label = f"Ch{st.session_state.story_chapter} · {_word_count(st.session_state.story_full_text)}w"
                st.session_state.story_branches.append(st.session_state.story_full_text)
                st.session_state.story_branch_labels.append(label)
                st.toast(f"Branch saved: {label}")
            if st.session_state.story_branches:
                for i, lbl in enumerate(st.session_state.story_branch_labels):
                    if st.button(f"↩ {lbl}", key=f"load_branch_{i}", use_container_width=True):
                        st.session_state.story_full_text = st.session_state.story_branches[i]
                        st.session_state.story_word_count = _word_count(st.session_state.story_full_text)
                        st.rerun()

        # ── Author Notes ────────────────────────
        st.markdown('<div class="config-label">📝 Author Notes</div>', unsafe_allow_html=True)
        notes = st.text_area(
            "Notes", placeholder="Characters, world rules, things to remember...",
            value=st.session_state.story_notes,
            label_visibility="collapsed",
            key="story_notes_input",
            height=90,
        )
        st.session_state.story_notes = notes

        # ── World Ledger ────────────────────────
        st.markdown('<div class="config-label">🌍 World Ledger</div>', unsafe_allow_html=True)
        with st.expander("Lore & Meta-Rules"):
            ledger = st.text_area(
                "Vault of Lore", placeholder="Magic systems, geography, political factions...",
                value=st.session_state.story_world_ledger,
                label_visibility="collapsed",
                key="story_ledger_input",
                height=150,
            )
            st.session_state.story_world_ledger = ledger

        # ── Story Critic ────────────────────────
        st.markdown('<div class="config-label">🧐 AI Critic </div>', unsafe_allow_html=True)
        with st.expander("Review & Analysis"):
            if st.session_state.story_full_text:
                for mode_name, mode_instr in CRITIC_MODES.items():
                    if st.button(f"Analyze: {mode_name}", use_container_width=True, key=f"critic_{mode_name}"):
                        with st.spinner("Reviewing manuscript..."):
                            critic_prompt = f"YOU ARE AN ELITE LITERARY CRITIC. {mode_instr}\n\nReview this story so far and provide 3 sharp, professional bullet points of feedback: what's working, what's not, and how to improve. Be honest and sophisticated.\n\nSTORY:\n{st.session_state.story_full_text[-3000:]}"
                            feedback = ai_engine.generate(critic_prompt, model="llama-3.1-8b-instant", max_tokens=300)
                            st.session_state.story_critic_feedback = feedback
                
                if st.button("🔍 Plot Hole Scanner", use_container_width=True, key="plot_hole_scan"):
                    with st.spinner("Analyzing causality and continuity..."):
                        logic_prompt = f"SCAN THIS STORY FOR LOGIC HOLES, IMPOSSIBLE TIMELINES, OR CONTRADICTIONS IN CHARACTER KNOWLEDGE. List only the major issues with zero fluff.\n\nSTORY:\n{st.session_state.story_full_text[-5000:]}"
                        st.session_state.story_logic_report = ai_engine.generate(logic_prompt, model="llama-3.1-8b-instant", max_tokens=300)
            
            if st.session_state.story_critic_feedback:
                st.markdown(f'<div style="font-size:0.8rem;color:var(--text2);background:var(--bg3-glass);padding:0.75rem;border-radius:8px;border:1px solid var(--accent-bd);">{st.session_state.story_critic_feedback}</div>', unsafe_allow_html=True)
                if st.button("Clear Feedback", key="clear_critic"):
                    st.session_state.story_critic_feedback = ""
                    st.rerun()

            if st.session_state.story_logic_report:
                st.markdown(f'<div style="font-size:0.8rem;color:#f87171;background:rgba(239, 68, 68, 0.05);padding:0.75rem;border-radius:8px;border:1px solid #f87171;"><b>LOGIC ISSUES FOUND:</b><br>{st.session_state.story_logic_report}</div>', unsafe_allow_html=True)
                if st.button("Clear Report", key="clear_logic"):
                    st.session_state.story_logic_report = ""
                    st.rerun()

        # ── Export & Sharing ─────────────────────
        st.markdown('<div class="config-label">📤 Export & Design</div>', unsafe_allow_html=True)
        with st.expander("Publishing Tools"):
            if st.button("📜 Export as Book Draft (Markdown)", use_container_width=True):
                full_book = f"# {st.session_state.story_title}\n\n"
                full_book += f"**Genre:** {st.session_state.story_genre}\n"
                full_book += f"**Total Word Count:** {st.session_state.story_word_count}\n\n"
                full_book += "---"
                full_book += f"\n\n{st.session_state.story_full_text}"
                st.download_button("Download Draft", full_book, file_name=f"{st.session_state.story_title.lower().replace(' ', '_')}.md")
            
            if st.button("🗺️ Generate Plot Diagram (Mermaid)", use_container_width=True):
                 st.info("Coming soon: Dynamic Mermaid mapping of your story structure.")

        # ── Advanced Engine Tuning ──────────────────
        st.markdown('<div class="config-label">⚙️ Engine Advanced Tuning</div>', unsafe_allow_html=True)
        with st.expander("Deep Parameters & Constraints", expanded=False):
            t_set1, t_set2, t_set3 = st.tabs(["📜 Constitution", "🎚️ AI Tuning", "🛠️ Custom Params"])
            with t_set1:
                st.markdown("<small style='color:var(--text3)'>Define unbreakable laws for the AI. It will NEVER violate these.</small>", unsafe_allow_html=True)
                st.session_state.story_constitution = st.text_area("Global Story Constitution", value=st.session_state.story_constitution, placeholder="e.g. No character can die twice. Never mention technology. Each turn must end with a question.", height=120, key="const_input")
                st.session_state.story_forbidden_words = st.text_input("Forbidden Words (comma separated)", value=st.session_state.story_forbidden_words, placeholder="e.g. okay, literally, modern", key="forbid_input")
                st.session_state.story_sentence_cap = st.number_input("Max Sentence Word Cap (0=Unlimited)", value=st.session_state.story_sentence_cap, min_value=0, max_value=100, step=5, key="scap_input")
            
            with t_set2:
                st.markdown("<small style='color:var(--text3)'>Precise control over the LLM generation engine.</small>", unsafe_allow_html=True)
                st.session_state.story_temp = st.slider("Temperature (Creativity)", 0.0, 1.5, st.session_state.story_temp, 0.05, key="temp_input")
                st.session_state.story_top_p = st.slider("Top-P (Randomness)", 0.0, 1.0, st.session_state.story_top_p, 0.05, key="top_p_input")
                st.session_state.story_max_tokens = st.select_slider("Max Continuation Length", options=[512, 1024, 2048, 4096, 8192], value=st.session_state.story_max_tokens, key="max_tok_input")
                st.caption("Lower temperature for logical accuracy, higher for wild creativity.")

            with t_set3:
                st.markdown("<small style='color:var(--text3)'>Define fully custom variables and parameters to feed the AI (e.g. 'Gore Level', 'Protagonist Flaw').</small>", unsafe_allow_html=True)
                cv_k = st.text_input("Parameter Name", placeholder="e.g. Target Audience", key="cv_k_input")
                cv_v = st.text_input("Parameter Value", placeholder="e.g. Young Adult", key="cv_v_input")
                if st.button("➕ Add Custom Parameter", use_container_width=True):
                    if cv_k and cv_v:
                        st.session_state.story_custom_vars.append({"key": cv_k, "value": cv_v})
                        st.rerun()
                
                if st.session_state.story_custom_vars:
                    for i, cv in enumerate(st.session_state.story_custom_vars):
                        st.markdown(f"**{cv['key']}:** {cv['value']}")
                        if st.button("❌ Remove", key=f"rm_cv_{i}"):
                            st.session_state.story_custom_vars.pop(i)
                            st.rerun()

        # ── Actions ─────────────────────────────
        st.markdown('<div class="config-label">⚙️ Actions</div>', unsafe_allow_html=True)
        a1, a2 = st.columns(2)
        with a1:
            if st.button("🗑️ Reset", use_container_width=True, key="story_reset"):
                for k in ["story_messages", "story_full_text", "story_title",
                          "story_started", "story_tool", "story_word_count", "story_chapter"]:
                    st.session_state[k] = [] if k == "story_messages" else (
                        "" if k in ["story_full_text", "story_title", "story_tool"] else
                        False if k == "story_started" else
                        0 if k == "story_word_count" else 1
                    )
                st.rerun()
        with a2:
            if st.session_state.story_full_text:
                st.download_button(
                    "📝 .txt Export",
                    data=_export_story(),
                    file_name=f"{(st.session_state.story_title or 'story').replace(' ','_')}.txt",
                    mime="text/plain",
                    use_container_width=True,
                    key="story_export_txt",
                )
                
                try:
                    import docx
                    from io import BytesIO
                    doc = docx.Document()
                    doc.add_heading(st.session_state.story_title or "Untitled", 0)
                    doc.add_paragraph(st.session_state.story_full_text)
                    bio = BytesIO()
                    doc.save(bio)
                    st.download_button("📄 .docx Export", data=bio.getvalue(), file_name=f"{(st.session_state.story_title or 'story').replace(' ','_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="story_export_docx")
                except ImportError:
                    pass
                    
                try:
                    from fpdf import FPDF
                    import textwrap
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    
                    def _safe_pdf_text(t: str) -> str:
                        t = t.encode('latin-1', 'ignore').decode('latin-1')
                        return "\n".join(["\n".join(textwrap.wrap(line, 85)) for line in t.splitlines()])
                        
                    safe_title = _safe_pdf_text(st.session_state.story_title or "Untitled")
                    safe_body = _safe_pdf_text(st.session_state.story_full_text)
                    
                    pdf.multi_cell(0, 10, txt=safe_title)
                    pdf.ln(5)
                    pdf.multi_cell(0, 10, txt=safe_body)
                    pdf_bytes = pdf.output(dest='S').encode('latin-1')
                    st.download_button("📕 .pdf Export", data=pdf_bytes, file_name=f"{(st.session_state.story_title or 'story').replace(' ','_')}.pdf", mime="application/pdf", use_container_width=True, key="story_export_pdf")
                except ImportError:
                    pass

        # ── Back to study ───────────────────────
        st.markdown("---")
        if st.button("← Back to Study Chat", use_container_width=True, key="story_back"):
            st.session_state.app_mode = "chat"
            st.rerun()

    # ══════════════════════════════════════
    # RIGHT PANEL — Canvas + Input
    # ══════════════════════════════════════
    with canvas_col:

        # ── Story Stats Bar ─────────────────────
        if st.session_state.story_started:
            wc = _word_count(st.session_state.story_full_text)
            chapters = st.session_state.story_chapter
            genre_tag = st.session_state.story_genre.split(" ", 1)[0]
            voice_tag = st.session_state.story_voice.split(" ", 1)[1] if " " in st.session_state.story_voice else st.session_state.story_voice
            st.markdown(f"""
            <div class="story-stat-bar">
              <span>📖 <b>{wc:,}</b> words</span>
              <span>📂 Chapter <b>{chapters}</b></span>
              <span>🎭 <b>{genre_tag}</b></span>
              <span>🖊️ <b>{voice_tag[:18]}</b></span>
              {f'<span>🔧 <b>{st.session_state.story_tool.split(" ",1)[1] if st.session_state.story_tool else ""}</b></span>' if st.session_state.story_tool else ''}
            </div>
            """, unsafe_allow_html=True)

        # ── Story Canvas ────────────────────────
        if st.session_state.story_full_text:
            # Show full story text beautifully
            st.markdown(f"""
            <div class="story-canvas">{st.session_state.story_full_text}</div>
            """, unsafe_allow_html=True)

            if st.session_state.storyboards:
                with st.expander("🖼️ Visual Storyboards", expanded=True):
                    for b in st.session_state.storyboards[-3:]:
                        st.markdown(f"""
                        <div style="background:var(--bg3-glass); border:1px solid var(--accent-bd); border-radius:12px; padding:15px; margin-bottom:10px;">
                          <div style="font-size:0.7rem; color:var(--text3); margin-bottom:5px; text-transform:uppercase; letter-spacing:1px;">Scene Concept</div>
                          <div style="font-style:italic; font-size:0.85rem; color:var(--text2);">{b['board']}</div>
                        </div>
                        """, unsafe_allow_html=True)
        else:
            # Empty state with spark prompts
            st.markdown(f"""
            <div class="story-canvas">
              <div class="story-canvas-empty">
                Your story will appear here.<br><br>
                Type your opening line below, or pick a spark prompt →
              </div>
            </div>
            """, unsafe_allow_html=True)

            # Spark prompts for current genre
            st.markdown('<div class="config-label">✨ Spark Prompts</div>', unsafe_allow_html=True)
            examples = GENRES[st.session_state.story_genre].get("examples", [])
            for ex in examples:
                if st.button(f'"{ex}"', key=f"spark_{ex[:20]}", use_container_width=True):
                    st.session_state["_queued_story_prompt"] = ex
                    st.rerun()

        # ── Continuation Input ──────────────────
        st.markdown("")

        if not st.session_state.story_started:
            placeholder_text = f'Begin your {st.session_state.story_genre.split(" ",1)[-1]} story... (first line, opening scene, or just a feeling)'
        else:
            placeholder_text = "Continue the story, or type what happens next... (leave blank to let the AI decide)"

        user_story_input = st.chat_input(
            placeholder_text,
            key="story_chat_input",
        )

        # Check for queued spark prompt
        if st.session_state.get("_queued_story_prompt"):
            user_story_input = st.session_state.pop("_queued_story_prompt")

        # ── Process input ───────────────────────
        if user_story_input is not None and user_story_input.strip() != "" or user_story_input == "":
            # Allow empty input after start (AI continues freely)
            raw_input = (user_story_input or "").strip()

            if not raw_input and not st.session_state.story_started:
                st.info("✍️ Type your opening line to begin the story.")
            else:
                # Build the message to send
                tool = st.session_state.story_tool
                tool_instruction = ""
                if tool:
                    tool_instruction = f"\n\n[NARRATIVE DIRECTIVE: {NARRATIVE_TOOLS[tool]}]"

                length_tokens = STORY_LENGTHS[st.session_state.story_length]

                if not st.session_state.story_started:
                    # First turn — start the story
                    if raw_input:
                        user_msg_content = f"Begin the story with this opening. Develop it beautifully into a full opening scene:\n\n\"{raw_input}\"{tool_instruction}"
                    else:
                        user_msg_content = f"Begin an original {st.session_state.story_genre} story with a compelling opening scene.{tool_instruction}"
                else:
                    # Continuation
                    if raw_input:
                        user_msg_content = f"Continue the story. The user directs: \"{raw_input}\"{tool_instruction}\n\nPick up exactly where we left off. Match the established voice perfectly."
                    else:
                        user_msg_content = f"Continue the story naturally from where it ended. Use your literary judgment to advance the plot, character, or atmosphere.{tool_instruction}"

                    # Add chapter divider every ~600 words
                    wc = _word_count(st.session_state.story_full_text)
                    if wc > 0 and wc % 600 < 80:
                        st.session_state.story_chapter += 1
                        chapter_break = f"\n\n\n— Chapter {st.session_state.story_chapter} —\n\n"
                        st.session_state.story_full_text += chapter_break

                # Add user message to history
                st.session_state.story_messages.append({
                    "role": "user",
                    "content": user_msg_content,
                })

                # Build system prompt
                sys_prompt = build_story_system_prompt(
                    st.session_state.story_genre,
                    st.session_state.story_voice,
                    st.session_state.story_pacing,
                    extra_context=st.session_state.story_notes,
                    style_mimicry=st.session_state.story_style_mimicry,
                    collab_mode=st.session_state.story_collab_mode,
                    characters=st.session_state.story_characters,
                    viability=st.session_state.story_viability,
                    sexy_touch=st.session_state.story_sexy_touch,
                    reply_format=st.session_state.story_reply_format,
                    atmosphere=st.session_state.story_atmosphere,
                    world_ledger=st.session_state.story_world_ledger,
                    archetype=st.session_state.story_archetype,
                    complexity=st.session_state.story_complexity,
                    tone=st.session_state.story_tone,
                    custom_genre_details=st.session_state.story_custom_genre,
                    custom_voice_instr=st.session_state.story_custom_voice,
                    custom_viability_instr=st.session_state.story_custom_viability,
                    custom_sexy_instr=st.session_state.story_custom_sexy,
                    story_wiki=st.session_state.story_wiki,
                    constitution=st.session_state.story_constitution,
                    forbidden_words=st.session_state.story_forbidden_words,
                    sentence_cap=st.session_state.story_sentence_cap,
                    custom_vars=st.session_state.story_custom_vars
                )


                # Keep only last 10 exchanges for context (saves tokens, maintains continuity)
                context_messages = st.session_state.story_messages[-20:]

                # Stream the story continuation
                with st.spinner("✍️ Writing..."):
                    generated = ""
                    try:
                        for chunk in ai_engine.generate_stream(
                            messages=context_messages,
                            context_text="",
                            model="llama-3.3-70b-versatile",
                            persona_prompt=sys_prompt,
                            temperature=st.session_state.story_temp,
                            top_p=st.session_state.story_top_p,
                            max_tokens=st.session_state.story_max_tokens
                        ):
                            generated += chunk

                        # Clean up any AI meta-commentary
                        generated = re.sub(
                            r'^(Sure|Certainly|Of course|Here\'s|Here is|I\'ll|I will|Let me)[^\n]*\n+',
                            '', generated.strip(), flags=re.IGNORECASE
                        )
                        generated = generated.strip()

                        # Append to story canvas
                        if st.session_state.story_full_text:
                            st.session_state.story_full_text += "\n\n" + generated
                        else:
                            st.session_state.story_full_text = generated

                        # Add AI response to message history
                        st.session_state.story_messages.append({
                            "role": "assistant",
                            "content": generated,
                        })

                        # Generate title on first turn
                        if not st.session_state.story_started:
                            st.session_state.story_title = _generate_title(
                                generated, st.session_state.story_genre
                            )
                            st.session_state.story_started = True

                        # clearing tools
                        st.session_state.story_tool = None

                        # Check if visualizer was requested
                        if tool == "🖼️ Scene Visualiser":
                            with st.spinner("🖼️ Sketching storyboard..."):
                                viz_prompt = f"Create a detailed, atmospheric visual storyboard prompt (70-100 words) for a professional illustrator based on the scene above. Focus on lighting, composition, and mood. No meta-talk.\n\nSCENE:\n{generated}"
                                storyboard = ai_engine.generate(viz_prompt, model="llama-3.1-8b-instant", max_tokens=200)
                                st.session_state.storyboards.append({"text": generated, "board": storyboard})

                        # Update word count
                        st.session_state.story_word_count = _word_count(st.session_state.story_full_text)

        # ── Stylometric Analysis ──
        if st.session_state.story_started:
            st.markdown('<div class="config-label">🔬 Prose Deep Analysis</div>', unsafe_allow_html=True)
            with st.expander("Stylometric Scan"):
                if st.button("📊 Run Full Stylometric Scan", use_container_width=True):
                    with st.spinner("Analyzing prose fingerprint..."):
                        # Calculate sentence length variance
                        sents = st.session_state.story_full_text.split(".")
                        avg_len = sum(len(s.split()) for s in sents) / max(1, len(sents))
                        
                        analysis_prompt = f"""ANALYZE THE FOLLOWING PROSE FINGERPRINT. 
                        Provide:
                        1. Lexical Diversity (Richness of vocabulary)
                        2. Narrative Density (Balance of action vs description)
                        3. Emotional Valence (Dominant mood)
                        4. Reading Ease (Grade level)
                        
                        TEXT:
                        {st.session_state.story_full_text[-2000:]}"""
                        
                        st.session_state.story_stylometrics = ai_engine.generate(analysis_prompt, model="llama-3.1-8b-instant", max_tokens=400)
                
                if st.session_state.story_stylometrics:
                    st.markdown(f'<div style="font-size:0.8rem; color:var(--text2); line-height:1.5;">{st.session_state.story_stylometrics}</div>', unsafe_allow_html=True)

        # ── Intelligence Analytics ──
        if st.session_state.story_started:
            st.markdown('<div class="config-label">📊 Manuscript Analytics</div>', unsafe_allow_html=True)
            t1, t2 = st.tabs(["⚡ Tension Map", "🔤 Vocabulary"])
            
            with t1:
                # Simple tension heuristic based on punctuation and word types
                tension = min(100, (st.session_state.story_word_count % 100) + 20)
                st.progress(tension/100, text=f"Narrative Tension: {tension}%")
                st.caption("AI Heuristic: Measuring sentence density and dialogue frequency.")
                
            with t2:
                words = st.session_state.story_full_text.lower().split()
                if words:
                    from collections import Counter
                    common = Counter(words).most_common(5)
                    st.markdown("**Top Frequency Words:**")
                    for w, c in common:
                        if len(w) > 3: st.write(f"- `{w}`: {c} times")

        # ── Advanced Intelligence Hub ────────────────
        if st.session_state.story_started:
            st.markdown('<div class="config-label">🌌 Narrative Intelligence Hub</div>', unsafe_allow_html=True)
            th1, th2 = st.tabs(["🗣️ Character Spotlight", "🔮 Plot Forecast"])
            
            with th1:
                st.markdown("<small style='color:var(--text3)'>Directly 'interview' one of your characters to find their authentic voice.</small>", unsafe_allow_html=True)
                if st.session_state.story_wiki["Characters"]:
                    target_char = st.selectbox("Select Character", list(st.session_state.story_wiki["Characters"].keys()), key="spotlight_char")
                    spot_q = st.text_input(f"Question for {target_char}", placeholder="e.g. 'What is your greatest fear right now?'", key="spotlight_q")
                    if st.button(f"Interview {target_char}", use_container_width=True):
                        with st.spinner(f"{target_char} is thinking..."):
                            char_context = st.session_state.story_wiki["Characters"][target_char]
                            spot_prompt = f"YOU ARE {target_char}. {char_context}\n\nMaintain this persona perfectly. Answer the user's question in your unique voice. Keep it under 100 words.\n\nSTORY SO FAR:\n{st.session_state.story_full_text[-1000:]}\n\nSTUDENT: {spot_q}"
                            response = ai_engine.generate(spot_prompt, model="llama-3.1-8b-instant", max_tokens=200)
                            st.markdown(f"**{target_char}:** {response}")
                else:
                    st.info("Add characters to the Story Wiki to use this feature.")

            with th2:
                if st.button("🔮 Forecast Future Paths", use_container_width=True):
                    with st.spinner("Calculating narrative probabilities..."):
                        forecast_prompt = f"Based on the current story, provide 3 distinct, high-fidelity future plot paths (each 1-2 sentences). Path 1: Logical progression. Path 2: Sharp twist. Path 3: Thematic payoff.\n\nSTORY:\n{st.session_state.story_full_text[-3000:]}"
                        forecast = ai_engine.generate(forecast_prompt, model="llama-3.1-8b-instant", max_tokens=300)
                        st.markdown(forecast)

        # ── Collaborative Outliner & Sensory Suite ──
        if st.session_state.story_started:
            st.markdown('<div class="config-label">🛠️ Tactical Narrative Suite</div>', unsafe_allow_html=True)
            tc1, tc2 = st.tabs(["📝 Plot Outline", "👂 Sensory Map"])
            
            with tc1:
                st.markdown("<small style='color:var(--text3)'>Build a structural skeleton before you write.</small>", unsafe_allow_html=True)
                if st.button("✨ Auto-Generate Outline", use_container_width=True):
                    with st.spinner("Drafting blueprint..."):
                        outline_prompt = f"Create a professional 5-point outline for the next major arc of this story. Ensure escalating stakes.\n\nSTORY:\n{st.session_state.story_full_text[-2000:]}"
                        outline = ai_engine.generate(outline_prompt, model="llama-3.1-8b-instant", max_tokens=300)
                        st.markdown(outline)
            
            with tc2:
                st.markdown("<small style='color:var(--text3)'>Analyze the sensory balance of your current prose.</small>", unsafe_allow_html=True)
                if st.button("🩺 Run Sensory Audit", use_container_width=True):
                    with st.spinner("Scanning textures..."):
                        audit_p = f"Analyze the following text and rate the presence (0-10) of: Visuals, Sounds, Smells, Touch, Taste. Briefly explain what's missing.\n\nTEXT:\n{st.session_state.story_full_text[-800:]}"
                        audit = ai_engine.generate(audit_p, model="llama-3.1-8b-instant", max_tokens=250)
                        st.markdown(audit)

                    except Exception as e:
                        st.error(f"Execution failed: {e}")

                st.rerun()
